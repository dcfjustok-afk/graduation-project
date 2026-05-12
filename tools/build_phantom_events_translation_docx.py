from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "output" / "Phantom Events-中文译文.md"
OUT = ROOT / "output" / "Phantom Events-中文译文.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_twips: int) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_twips))
    tbl_w.set(qn("w:type"), "dxa")
    jc = tbl_pr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        tbl_pr.append(jc)
    jc.set(qn("w:val"), "left")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_doc_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.7)
    section.right_margin = Cm(2.7)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "SimSun"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color in [
        ("Title", 20, "1F4E79"),
        ("Heading 1", 15, "1F4E79"),
        ("Heading 2", 13, "2F5F8F"),
        ("Heading 3", 11.5, "333333"),
    ]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "SimSun"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.size = Pt(10.5)

    for section in doc.sections:
        header = section.header
        p = header.paragraphs[0]
        p.text = "Phantom Events 中文译文"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if p.runs:
            p.runs[0].font.name = "Microsoft YaHei"
            p.runs[0]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
            p.runs[0].font.size = Pt(9)
            p.runs[0].font.color.rgb = RGBColor(90, 90, 90)
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        fp.add_run("第 ")
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), "PAGE")
        fp._p.append(fld)
        fp.add_run(" 页")


def add_rich_runs(paragraph, text: str, font_size: float | None = None, mono: bool = False) -> None:
    # Keep inline code/math markers as literal text but render code spans in monospace.
    text = clean_display_text(text)
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            run.font.size = Pt(font_size or 9.5)
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.name = "Microsoft YaHei"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
            if font_size:
                run.font.size = Pt(font_size)
        else:
            run = paragraph.add_run(part)
            if mono:
                run.font.name = "Consolas"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            else:
                run.font.name = "SimSun"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            if font_size:
                run.font.size = Pt(font_size)


def clean_display_text(text: str) -> str:
    """Make common LaTeX fragments readable in DOCX while preserving identifiers."""
    text = re.sub(r"\\\((.*?)\\\)", r"\1", text)
    text = re.sub(r"\\bm\{\\mathcal\{M\}\}_\{([^}]+)\}", r"M_\1", text)
    text = re.sub(r"\\texttt\{([^}]+)\}", r"\1", text)
    text = re.sub(r"\\textit\{([^}]+)\}", r"\1", text)
    text = re.sub(r"\\textsc\{([^}]+)\}", r"\1", text)
    text = re.sub(r"\\emph\{([^}]+)\}", r"\1", text)
    text = text.replace("\\_", "_")
    text = text.replace("\\ldots{}", "...")
    return text


def add_paragraph(doc: Document, text: str, style: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    add_rich_runs(p, text)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text.strip("*"))
    run.bold = True
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(70, 70, 70)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(5)


def add_reference_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.first_line_indent = Cm(-0.6)
    p.paragraph_format.space_after = Pt(3)
    add_rich_runs(p, text, font_size=9)


def split_table_row(line: str) -> list[str]:
    body = line.strip().strip("|")
    return [cell.strip() for cell in body.split("|")]


def add_markdown_table(doc: Document, lines: list[str]) -> None:
    rows = [split_table_row(line) for line in lines if line.strip().startswith("|")]
    rows = [row for row in rows if not all(re.fullmatch(r":?-{3,}:?", c or "") for c in row)]
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    rows = [r + [""] * (ncols - len(r)) for r in rows]

    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(table, 8850)
    for i, row in enumerate(rows):
        tr = table.rows[i]
        if i == 0:
            set_repeat_table_header(tr)
        for j, value in enumerate(row):
            cell = tr.cells[j]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            if i == 0:
                set_cell_shading(cell, "EAF2F8")
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_rich_runs(p, value, font_size=9)
            for run in p.runs:
                if i == 0:
                    run.bold = True
                    run.font.name = "Microsoft YaHei"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    doc.add_paragraph()


def add_code_block(doc: Document, lines: list[str]) -> None:
    for raw in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.35)
        p.paragraph_format.right_indent = Cm(0.2)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        p_pr = p._p.get_or_add_pPr()
        shd = p_pr.find(qn("w:shd"))
        if shd is None:
            shd = OxmlElement("w:shd")
            p_pr.append(shd)
        shd.set(qn("w:fill"), "F7F7F7")
        run = p.add_run(raw or " ")
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(8)
    doc.add_paragraph()


def build_docx() -> None:
    text = SRC.read_text(encoding="utf-8")
    doc = Document()
    set_doc_defaults(doc)

    in_code = False
    code_lines: list[str] = []
    table_lines: list[str] = []

    def flush_table():
        nonlocal table_lines
        if table_lines:
            add_markdown_table(doc, table_lines)
            table_lines = []

    def flush_code():
        nonlocal code_lines
        if code_lines:
            add_code_block(doc, code_lines)
            code_lines = []

    for raw in text.splitlines():
        line = raw.rstrip()
        if line.startswith("```"):
            if in_code:
                flush_code()
                in_code = False
            else:
                flush_table()
                in_code = True
            continue

        if in_code:
            code_lines.append(line)
            continue

        if line.strip().startswith("|"):
            table_lines.append(line)
            continue
        flush_table()

        stripped = line.strip()
        if not stripped:
            continue

        if stripped.startswith("# "):
            p = doc.add_paragraph(style="Title")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_rich_runs(p, stripped[2:], font_size=20)
            continue
        if stripped.startswith("## "):
            doc.add_page_break()
            add_paragraph(doc, stripped[3:], style="Heading 1")
            continue
        if stripped.startswith("### "):
            add_paragraph(doc, stripped[4:], style="Heading 1")
            continue
        if stripped.startswith("#### "):
            add_paragraph(doc, stripped[5:], style="Heading 2")
            continue
        if stripped.startswith("##### "):
            add_paragraph(doc, stripped[6:], style="Heading 3")
            continue
        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_rich_runs(p, stripped[2:])
            continue
        if re.match(r"^\[\d+\]\s", stripped):
            add_reference_paragraph(doc, stripped)
            continue
        if stripped.startswith("**图 ") or stripped.startswith("**表 ") or stripped.startswith("**算法 "):
            add_caption(doc, stripped)
            continue
        if stripped.startswith("$$") and stripped.endswith("$$") and len(stripped) > 4:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_rich_runs(p, stripped, font_size=9.5, mono=True)
            continue
        if stripped in {"### 致谢", "### 参考文献"}:
            add_paragraph(doc, stripped[4:], style="Heading 1")
            continue

        add_paragraph(doc, stripped)

    flush_table()
    flush_code()
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_docx()
