from __future__ import annotations

from pathlib import Path

from docx import Document


DOCX_PATH = Path("output/Graduation-thesis.docx")


def is_heading(text: str) -> bool:
    text = text.strip()
    if not text:
        return False
    if text in {"摘要", "Abstract", "致谢", "参考文献"}:
        return True
    if text.startswith("关键词：") or text.startswith("Keywords:"):
        return True
    if text.startswith("第 ") and "章" in text:
        return True
    parts = text.split(maxsplit=1)
    if not parts:
        return False
    marker = parts[0]
    if marker.count(".") >= 1 and all(part.isdigit() for part in marker.split(".")):
        return True
    return False


def main() -> None:
    doc = Document(DOCX_PATH)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    full = "\n".join(paragraphs)

    print(f"file={DOCX_PATH.resolve()}")
    print(f"paragraphs={len(paragraphs)}")
    print(f"chars={sum(len(p) for p in paragraphs)}")
    print(f"tables={len(doc.tables)}")
    print(f"question_count={full.count('?')}")
    print(f"section_summary_count={full.count('本节关键信息摘要')}")
    print()

    current = "__front__"
    body: list[str] = []
    sections: list[tuple[str, list[str]]] = []
    for text in paragraphs:
        if is_heading(text):
            if body or current != "__front__":
                sections.append((current, body))
            current = text
            body = []
        else:
            body.append(text)
    sections.append((current, body))

    for title, paras in sections:
        chars = sum(len(p) for p in paras)
        if chars:
            print(f"{title}\tparas={len(paras)}\tchars={chars}")

    print("\nKEYS")
    keys = [
        "LogRegistry",
        "LOGGER_ROLE",
        "provider.getCode",
        "contract_address",
        "expectedHash",
        "actualHash",
        "onChainHash",
        "hash_mismatch",
        "107.03",
        "9.33",
        "3067.77",
        "15659.44",
        "35032.13",
        "auditStatus",
        "alertGenerated",
    ]
    for key in keys:
        print(f"{key}\t{full.count(key)}")


if __name__ == "__main__":
    main()
