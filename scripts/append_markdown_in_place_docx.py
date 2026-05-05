from copy import deepcopy
from pathlib import Path
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_TARGET = ROOT / "output" / "Graduation-thesis.docx"


def set_font(run, size=10.5, bold=False):
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold


def set_normal_style(document):
    style = document.styles["Normal"]
    style.font.name = "宋体"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(10.5)


def add_heading(document, text, level):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.5
    if level == 1:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    size = 16 if level == 1 else 15 if level == 2 else 12
    run = paragraph.add_run(text)
    set_font(run, size=size, bold=True)


def add_paragraph(document, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Pt(21)
    paragraph.paragraph_format.line_spacing = 1.5
    run = paragraph.add_run(text)
    set_font(run)


def add_bullet(document, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Pt(18)
    paragraph.paragraph_format.first_line_indent = Pt(-18)
    paragraph.paragraph_format.line_spacing = 1.5
    run = paragraph.add_run("• " + text)
    set_font(run)


def add_numbered(document, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Pt(18)
    paragraph.paragraph_format.first_line_indent = Pt(-18)
    paragraph.paragraph_format.line_spacing = 1.5
    run = paragraph.add_run(text)
    set_font(run)


def add_markdown(document, source):
    text = source.read_text(encoding="utf-8")
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("### "):
            add_heading(document, line[4:], 3)
        elif line.startswith("## "):
            add_heading(document, line[3:], 2)
        elif line.startswith("# "):
            add_heading(document, line[2:], 1)
        elif line.startswith("- "):
            add_bullet(document, line[2:])
        elif line[:2].isdigit() and line[2:4] == ". ":
            add_numbered(document, line)
        elif line[:1].isdigit() and line[1:3] == ". ":
            add_numbered(document, line)
        else:
            add_paragraph(document, line)


def build_fragment(source):
    fragment = Document()
    set_normal_style(fragment)
    fragment.add_page_break()
    add_markdown(fragment, source)
    return [child for child in fragment._body._element if not child.tag.endswith("sectPr")]


def append_markdown(source, target):
    document = Document(target)
    body = document._body._element
    sect_pr = None
    for child in body:
        if child.tag.endswith("sectPr"):
            sect_pr = child
            break
    for child in build_fragment(source):
        if sect_pr is None:
            body.append(deepcopy(child))
        else:
            sect_pr.addprevious(deepcopy(child))
    document.save(target)


if __name__ == "__main__":
    if len(sys.argv) not in (2, 3):
        raise SystemExit("usage: python scripts/append_markdown_in_place_docx.py <markdown-path> [target-docx]")
    source_path = Path(sys.argv[1])
    if not source_path.is_absolute():
        source_path = ROOT / source_path
    target_path = Path(sys.argv[2]) if len(sys.argv) == 3 else DEFAULT_TARGET
    if not target_path.is_absolute():
        target_path = ROOT / target_path
    append_markdown(source_path, target_path)
    print(target_path)
