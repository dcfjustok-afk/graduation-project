from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "output" / "stage1-sections-3.1-to-3.2.md"
TARGET = ROOT / "output" / "Graduation-thesis.docx"


def set_font(run, size=10.5, bold=False):
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold


def is_effectively_empty(document):
    return not any(paragraph.text.strip() for paragraph in document.paragraphs) and not document.tables


def clear_document(document):
    body = document._body._element
    for child in list(body):
        body.remove(child)


def add_heading(document, text, level):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.5
    if level == 1:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        size = 16
    elif level == 2:
        size = 15
    else:
        size = 12
    run = paragraph.add_run(text)
    set_font(run, size=size, bold=True)


def add_paragraph(document, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Pt(21)
    paragraph.paragraph_format.line_spacing = 1.5
    run = paragraph.add_run(text)
    set_font(run)


def add_bullet(document, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Pt(18)
    paragraph.paragraph_format.first_line_indent = Pt(-18)
    paragraph.paragraph_format.line_spacing = 1.5
    run = paragraph.add_run("• " + text)
    set_font(run)


def add_numbered(document, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Pt(18)
    paragraph.paragraph_format.first_line_indent = Pt(-18)
    paragraph.paragraph_format.line_spacing = 1.5
    run = paragraph.add_run(text)
    set_font(run)


def add_table_placeholder(document, text):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.line_spacing = 1.5
    run = paragraph.add_run(text)
    set_font(run, bold=True)


def add_simple_field_table(document):
    rows = [
        ("字段名称", "字段含义", "主要作用"),
        ("taskId", "任务标识", "归集同一任务下的日志，支持链上查询和审计回溯。"),
        ("sourceType", "来源类型", "描述日志提交来源，如 Agent 采集。"),
        ("sourcePath", "来源路径", "记录日志文件路径或来源位置。"),
        ("logContent", "日志原文", "用于链下保存、哈希计算和审计重算。"),
        ("logLevel", "日志级别", "用于区分 INFO、WARN、ERROR 等日志级别。"),
        ("collectedAt", "采集时间", "记录日志进入系统的时间。"),
    ]
    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, value in enumerate(rows[0]):
        table.rows[0].cells[idx].text = value
    for row in rows[1:]:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_font(run, bold=(row == table.rows[0]))


def add_chain_mapping_table(document):
    rows = [
        ("链下字段", "链上字段", "映射说明"),
        ("logs.id", "-", "链下日志主键，通过 log_hash_records.log_id 关联存证记录。"),
        ("logs.task_id", "taskId", "任务标识，用于任务维度归集和链上查询。"),
        ("log_hash_records.log_hash", "logHash", "日志哈希，是完整性校验的核心依据。"),
        ("log_hash_records.contract_address", "合约地址", "记录实际写入的 LogRegistry 地址，支持历史回溯审计。"),
        ("log_hash_records.transaction_hash", "交易哈希", "记录链上写入交易，便于追踪存证过程。"),
        ("log_hash_records.block_number", "区块号", "记录交易所在区块，辅助说明链上确认位置。"),
        ("log_hash_records.on_chain_status", "-", "记录链上写入状态或失败原因。"),
        ("-", "createdAt", "链上写入时的区块时间戳。"),
        ("-", "submitter", "执行链上写入的钱包地址。"),
    ]
    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, value in enumerate(rows[0]):
        table.rows[0].cells[idx].text = value
    for row in rows[1:]:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_font(run, bold=(row == table.rows[0]))


def write_markdown_to_docx():
    source_text = SOURCE.read_text(encoding="utf-8")
    if TARGET.exists():
        existing_document = Document(TARGET)
        document = Document() if is_effectively_empty(existing_document) else existing_document
    else:
        document = Document()

    style = document.styles["Normal"]
    style.font.name = "宋体"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(10.5)

    if not is_effectively_empty(document):
        document.add_page_break()

    for raw_line in source_text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("## "):
            add_heading(document, line[3:], 2)
        elif line.startswith("# "):
            add_heading(document, line[2:], 1)
        elif line.startswith("- "):
            add_bullet(document, line[2:])
        elif line.startswith("图 ") or line.startswith("图"):
            add_table_placeholder(document, line)
        elif line == "表 3-1 日志标准化封装字段说明":
            add_table_placeholder(document, line)
            add_simple_field_table(document)
        elif line == "表 3-1 链上链下字段映射关系":
            add_table_placeholder(document, line)
            add_chain_mapping_table(document)
        elif line[:2].isdigit() and line[2:4] == ". ":
            add_numbered(document, line)
        elif line[:1].isdigit() and line[1:3] == ". ":
            add_numbered(document, line)
        else:
            add_paragraph(document, line)

    document.save(TARGET)


if __name__ == "__main__":
    write_markdown_to_docx()
    print(TARGET)
