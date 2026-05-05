from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.shared import Pt


DOCX_PATH = Path("output/Graduation-thesis.docx")
PROMPT_PATH = Path("image/论文新增插图AI生成提示词-无标题版.md")
REPORT_PATH = Path("output/论文新增图片位置修改报告.md")


FIGURES = [
    {
        "caption": "图 2-1 区块链存证基本结构示意图",
        "filename": "fig-2-1-blockchain-evidence-structure.png",
        "section": "2.1 区块链与链上存证技术",
        "anchor_type": "section_end",
        "transition": "为进一步说明链上存证与链下数据之间的关系，图 2-1 对区块链存证的基本结构进行概括。该图用于说明业务数据并不必然完整上链，而是可以通过摘要、交易和区块记录形成可追溯证据。",
        "content": [
            "横向 16:9，白色背景，正式论文技术图风格。",
            "左侧为“链下业务数据”，包含日志原文、数据库记录、业务标识。",
            "中间为“哈希摘要生成”和“交易提交”，展示 SHA-256 摘要进入交易。",
            "右侧为连续区块，区块之间用前一区块哈希连接，标注交易摘要、时间戳、区块哈希。",
            "底部展示“后续审计校验”，从链下数据重新计算摘要并与链上记录比对。",
        ],
        "prompt": "生成一张本科软件工程毕业论文使用的区块链存证基本结构示意图。画面为横向 16:9，白色背景，正式、清晰、学术化，不要在图片内部写“图 2-1”或任何论文图题。左侧绘制链下业务数据区域，包含日志原文、数据库记录、业务标识等卡片；中间绘制 SHA-256 哈希摘要生成和交易提交过程；右侧绘制 3 至 4 个相连区块，每个区块内部包含交易摘要、时间戳、区块哈希、前一区块哈希等简化字段；底部绘制后续审计校验过程，从链下数据重新计算摘要并与链上记录比对。颜色以蓝色、青色、绿色为主，少量橙色用于校验节点。文字使用简洁中文标签，文字清晰无错别字。不要出现比特币 Logo、公司 Logo、水印、二维码、人物、卡通装饰、复杂渐变背景。输出 PNG，建议 1920x1080，保存为 image/fig-2-1-blockchain-evidence-structure.png。",
    },
    {
        "caption": "图 2-2 哈希完整性校验原理图",
        "filename": "fig-2-2-hash-integrity-verification.png",
        "section": "2.3 哈希摘要与日志完整性校验技术",
        "anchor_type": "section_end",
        "transition": "哈希摘要用于完整性校验时，重点并不在于还原日志原文，而在于比较同一输入在不同时刻得到的摘要是否一致。图 2-2 展示了日志内容变化导致哈希摘要变化的基本原理。",
        "content": [
            "上下两条对比流程：原始日志与修改后日志。",
            "两条流程都经过 SHA-256 计算节点，分别得到哈希摘要 A 与哈希摘要 B。",
            "右侧使用对比符号展示摘要不一致，并标注“内容变化可被检测”。",
            "突出单向性、固定长度摘要、雪崩效应、完整性校验。",
        ],
        "prompt": "生成一张本科毕业论文使用的哈希完整性校验原理图。画面为横向 16:9，白色背景，简洁学术风格，不要在图片内部写“图 2-2”或任何论文图题。上方流程为“原始日志内容 -> SHA-256 哈希计算 -> 哈希摘要 A”；下方流程为“修改后日志内容 -> SHA-256 哈希计算 -> 哈希摘要 B”。右侧用清晰的对比符号展示“哈希摘要 A 不等于哈希摘要 B”，并标注“日志内容变化可被检测”。图中可以用三个小标签说明“单向性、固定长度摘要、雪崩效应”，但不要展开过多密码学理论。颜色以蓝色和橙色为主，原始流程使用蓝色，修改后流程使用橙色或红色提示差异。文字必须清晰，不要错别字，不要使用英文长句，不要出现水印、Logo、人物或卡通装饰。输出 PNG，建议 1920x1080，保存为 image/fig-2-2-hash-integrity-verification.png。",
    },
    {
        "caption": "图 3-4 LogRegistry 合约数据结构关系图",
        "filename": "fig-3-4-logregistry-data-structure.png",
        "section": "3.3.1 合约数据结构设计",
        "anchor_type": "section_end",
        "transition": "为直观展示链上记录的组织方式，图 3-4 给出了 LogRecord 结构体、records 数组和 taskIdToRecordIds 映射之间的关系。",
        "content": [
            "中心为 LogRecord 结构体，字段 taskId、logHash、createdAt、submitter。",
            "左侧或上方展示 records 数组，按 recordId 顺序保存 LogRecord。",
            "右侧展示 taskIdToRecordIds 映射，一个 taskId 对应多个 recordId。",
            "展示 storeLog 写入和 getLogsByTaskId 查询方向。",
        ],
        "prompt": "生成一张本科软件工程毕业论文使用的智能合约数据结构关系图。画面为横向 16:9，白色背景，正式工程图风格，不要在图片内部写“图 3-4”或任何论文图题。中心绘制 LogRecord 结构体卡片，字段包括 taskId、logHash、createdAt、submitter；左侧绘制 records 数组，展示 recordId 0、recordId 1、recordId 2 等编号顺序保存 LogRecord；右侧绘制 taskIdToRecordIds 映射，展示同一个 taskId 可以对应多个 recordId。用箭头表示 storeLog 写入 records 和 taskIdToRecordIds，用另一组箭头表示 getLog 按编号查询、getLogsByTaskId 按任务查询。图中只展示论文中真实存在的合约字段和查询能力，不要加入余额、转账、NFT、Token 等无关内容。配色使用绿色表示链上合约，蓝色表示查询，橙色表示写入。输出 PNG，建议 1920x1080，保存为 image/fig-3-4-logregistry-data-structure.png。",
    },
    {
        "caption": "图 3-5 历史合约地址回溯审计策略图",
        "filename": "fig-3-5-contract-address-tracing-audit.png",
        "section": "3.4.2 历史合约地址回溯审计策略",
        "anchor_type": "section_end",
        "transition": "该策略的关键在于审计阶段不只依赖当前环境变量中的合约地址，而是优先回到日志写入时记录的 contract_address。图 3-5 展示了历史合约地址回溯的审计路径。",
        "content": [
            "左侧展示日志提交阶段：写入 LogRegistry，同时保存 contract_address、transaction_hash、block_number。",
            "右侧展示审计阶段：读取 log_hash_records.contract_address，连接历史合约地址查询 onChainHash。",
            "突出避免当前配置地址与历史写入地址不一致造成误判。",
        ],
        "prompt": "生成一张本科毕业论文使用的历史合约地址回溯审计策略图。画面为横向 16:9，白色背景，正式流程图风格，不要在图片内部写“图 3-5”或任何论文图题。左侧区域标题为“日志提交阶段”，展示 Server 调用 LogRegistry.storeLog(taskId, logHash)，并将 contract_address、transaction_hash、block_number、on_chain_status 写入 log_hash_records。右侧区域标题为“审计阶段”，展示审计服务读取 log_hash_records.contract_address，连接当时写入的 LogRegistry 合约地址，再查询 onChainHash。图中加入一个对比提示：如果只读取当前环境变量合约地址，可能因 Hardhat 本地链重启或多次部署造成历史日志误判；使用 contract_address 回溯可以定位历史写入合约。不要加入不存在的跨链、联盟链节点、公开链浏览器等功能。配色以蓝色表示后端，绿色表示合约，紫色表示历史地址，橙色表示风险提示。输出 PNG，建议 1920x1080，保存为 image/fig-3-5-contract-address-tracing-audit.png。",
    },
    {
        "caption": "图 3-6 合约代码存在性校验流程图",
        "filename": "fig-3-6-contract-code-existence-check.png",
        "section": "3.4.3 合约代码存在性校验机制",
        "anchor_type": "section_end",
        "transition": "除记录历史合约地址外，系统还需要确认目标地址上确实存在合约字节码。图 3-6 展示了基于 provider.getCode(address) 的前置校验流程。",
        "content": [
            "流程：读取 contract_address -> provider.getCode(address) -> 判断字节码是否为空。",
            "存在字节码：继续 storeLog 或 getLog/getLogsByTaskId。",
            "不存在字节码：不继续链上读写，记录 pending 或异常说明。",
            "突出本地链重启后空地址误判问题。",
        ],
        "prompt": "生成一张本科毕业论文使用的合约代码存在性校验流程图。画面为横向 16:9，白色背景，正式技术流程图风格，不要在图片内部写“图 3-6”或任何论文图题。流程节点依次为：读取 contract_address；调用 provider.getCode(address)；判断返回字节码是否为空；如果存在合约代码，则继续执行 storeLog 或 getLog/getLogsByTaskId；如果不存在合约代码，则停止链上读写，并将审计结果标记为 pending 或记录异常说明。图中需要明确说明该机制用于避免 Hardhat 本地链重启后空地址误判。使用菱形表示判断节点，绿色分支表示可继续调用，黄色或红色分支表示合约不可用。不要加入论文未实现的自动部署、自动修复、区块链浏览器查询等功能。输出 PNG，建议 1920x1080，保存为 image/fig-3-6-contract-code-existence-check.png。",
    },
    {
        "caption": "图 4-3 核心数据库 ER 关系图",
        "filename": "fig-4-3-database-er-relationship.png",
        "section": "4.7 数据库设计",
        "anchor_type": "after_table",
        "table_header": ["表名", "主要用途", "关键字段"],
        "transition": "表 4-1 从字段角度概括了核心数据表。为进一步说明表之间的关联关系，图 4-3 展示了 logs、log_hash_records、audit_records、alerts 和 agent_states 的主要联系。",
        "content": [
            "五张表：logs、log_hash_records、audit_records、alerts、agent_states。",
            "重点关系：logs.id -> log_hash_records.log_id；log_hash_records.id -> audit_records.log_hash_record_id。",
            "alerts.related_log_id -> logs.id；alerts.related_audit_id -> audit_records.id。",
            "agent_states 独立记录采集端状态，与日志采集来源相关。",
        ],
        "prompt": "生成一张本科软件工程毕业论文使用的核心数据库 ER 关系图。画面为横向 16:9，白色背景，正式数据库设计图风格，不要在图片内部写“图 4-3”或任何论文图题。图中包含五张表：logs、log_hash_records、audit_records、alerts、agent_states。每张表用矩形表结构表示，列出核心字段。logs 表包含 id、task_id、source_type、source_path、log_content、log_level、collected_at、status；log_hash_records 表包含 id、log_id、task_id、log_hash、contract_address、transaction_hash、block_number、on_chain_status；audit_records 表包含 id、log_id、log_hash_record_id、audit_status、expected_hash、actual_hash、audited_at；alerts 表包含 id、alert_type、related_log_id、related_audit_id、severity、status；agent_states 表包含 id、source_path、last_offset、last_heartbeat_at、last_sync_at、status。用连线展示 logs.id 到 log_hash_records.log_id，log_hash_records.id 到 audit_records.log_hash_record_id，logs.id 到 alerts.related_log_id，audit_records.id 到 alerts.related_audit_id。agent_states 可以单独放置，并标注用于记录 Agent 状态。不要加入论文中不存在的数据表或字段。输出 PNG，建议 1920x1080，保存为 image/fig-4-3-database-er-relationship.png。",
    },
    {
        "caption": "图 5-1 Agent 增量采集与偏移量持久化流程图",
        "filename": "fig-5-1-agent-incremental-offset-flow.png",
        "section": "5.2.1 增量读取与偏移量持久化实现",
        "anchor_type": "section_end",
        "transition": "Agent 增量采集的核心在于读取新增内容后再更新偏移量，避免提交失败时丢失日志。图 5-1 展示了增量读取、提交确认与偏移量持久化之间的关系。",
        "content": [
            "流程：读取 sourcePath -> 获取 last_offset -> 读取新增日志 -> logCollector 封装。",
            "提交 Server 成功后更新 offsetStore。",
            "提交失败进入 retryQueue，暂不更新偏移量。",
            "状态同步写入 agent_states。",
        ],
        "prompt": "生成一张本科毕业论文使用的 Agent 增量采集与偏移量持久化流程图。画面为横向 16:9，白色背景，正式软件工程流程图风格，不要在图片内部写“图 5-1”或任何论文图题。流程从左到右依次展示：读取 sourcePath；从 offsetStore 获取 last_offset；fileReader 从偏移量之后读取新增日志；logCollector 封装 taskId、sourceType、sourcePath、logContent、logLevel、collectedAt；提交 Server；如果提交成功，则更新 offsetStore 中的 last_offset，并同步 agent_states；如果提交失败，则进入 retryQueue，等待下一轮重试，且暂不更新偏移量。图中要突出“成功后更新偏移量、失败时保留待提交日志”的原则。不要加入消息队列 Kafka、Redis、容器平台等论文未使用技术。输出 PNG，建议 1920x1080，保存为 image/fig-5-1-agent-incremental-offset-flow.png。",
    },
    {
        "caption": "图 5-2 后端分层架构图",
        "filename": "fig-5-2-backend-layered-architecture.png",
        "section": "5.3.1 后端分层架构与核心接口实现",
        "anchor_type": "section_end",
        "transition": "后端分层结构使接口处理、业务编排、数据库访问和链上交互保持边界清晰。图 5-2 展示了 Express 后端各层与外部模块之间的关系。",
        "content": [
            "外部调用方：Agent、Web 前端、实验脚本。",
            "接口：POST /api/logs、POST /api/audits/run、GET /api/alerts、GET /api/logs。",
            "后端层次：routes、controllers、services、repositories、db、blockchain。",
            "外部资源：SQLite、LogRegistry。",
        ],
        "prompt": "生成一张本科软件工程毕业论文使用的后端分层架构图。画面为横向 16:9，白色背景，正式系统实现图风格，不要在图片内部写“图 5-2”或任何论文图题。左侧展示外部调用方：Agent、Web 前端、性能测试脚本。中间展示 Express 后端分层，从上到下或从左到右依次为 routes、controllers、services、repositories、db、blockchain。routes 层标注 POST /api/logs、GET /api/logs、POST /api/audits/run、GET /api/alerts；controllers 层标注参数读取、校验、响应封装；services 层标注日志服务、区块链服务、审计执行服务、告警服务、总览服务；repositories/db 层连接 SQLite；blockchain 层连接 LogRegistry。右侧展示 SQLite 和 LogRegistry 两个外部资源。不要加入论文未使用的 MySQL、Redis、Kafka、Docker、Kubernetes 等技术。颜色清晰、文字简洁，适合插入 Word。输出 PNG，建议 1920x1080，保存为 image/fig-5-2-backend-layered-architecture.png。",
    },
    {
        "caption": "图 5-3 篡改检测实验闭环示意图",
        "filename": "fig-5-3-tamper-detection-loop.png",
        "section": "5.6.3 篡改检测实验与结果分析",
        "anchor_type": "section_end",
        "transition": "为说明篡改检测实验中审计状态和告警生成之间的关系，图 5-3 展示了从日志修改到 failed 审计记录和 hash_mismatch 告警生成的闭环过程。",
        "content": [
            "正常提交：logs、log_hash_records、LogRegistry。",
            "修改日志内容。",
            "审计重算 actualHash，与 expectedHash、onChainHash 比对。",
            "auditStatus=failed，alertGenerated=true，生成 hash_mismatch 告警。",
        ],
        "prompt": "生成一张本科毕业论文使用的篡改检测实验闭环示意图。画面为横向 16:9，白色背景，正式实验流程图风格，不要在图片内部写“图 5-3”或任何论文图题。流程依次展示：正常日志提交；日志原文写入 logs 表；log_hash 写入 log_hash_records；logHash 写入 LogRegistry；随后日志内容被修改；审计阶段重新计算 actualHash；与 expectedHash、onChainHash 进行三方比对；比对不一致后生成 auditStatus=failed；同时 alertGenerated=true，并写入 hash_mismatch 告警。图中突出“日志内容被修改 -> 哈希不一致 -> failed -> 告警生成”的闭环。不要加入机器学习异常检测、自动修复、邮件通知等论文未实现功能。输出 PNG，建议 1920x1080，保存为 image/fig-5-3-tamper-detection-loop.png。",
    },
]


EXISTING_CAPTION_RENAMES = {
    "图 5-1 系统前端总览页面": "图 5-4 系统前端总览页面",
    "图 5-2 审计管理与告警管理页面": "图 5-5 审计管理与告警管理页面",
    "图 5-3 批量审计耗时对比图": "图 5-6 批量审计耗时对比图",
}


TEXT_REPLACEMENTS = {
    "页面通过 GET /api/logs 读取日志列表，并结合审计记录映射日志状态，展示任务标识、日志级别、来源路径、采集时间和当前状态等信息。":
        "页面通过 GET /api/logs 读取日志列表，并结合审计记录映射日志状态，展示任务标识、日志级别、来源路径、采集时间和当前状态等信息。系统总览页面效果如图 5-4 所示。",
    "对于 hash_mismatch 类型告警，页面能够提示其来源于日志哈希不一致，并通过关联日志和审计记录辅助用户追踪异常原因。前端数据处理主要集中在 API 客户端、dataService 和 mappers 中，使页面组件只关注展示逻辑，而将接口字段转换和状态映射集中处理。":
        "对于 hash_mismatch 类型告警，页面能够提示其来源于日志哈希不一致，并通过关联日志和审计记录辅助用户追踪异常原因。前端数据处理主要集中在 API 客户端、dataService 和 mappers 中，使页面组件只关注展示逻辑，而将接口字段转换和状态映射集中处理。审计管理与告警管理页面效果如图 5-5 所示。",
    "从结果分析看，批量审计耗时由数据库读取、哈希重算、链上记录查询、比对判断和审计结果写入共同构成。随着审计数据规模从 100 条增加到 1000 条，总耗时同步上升，说明当前原型系统仍存在批量审计优化空间。但三组数据 5 轮均执行成功，表明系统在实验规模下能够稳定完成审计闭环。":
        "从结果分析看，批量审计耗时由数据库读取、哈希重算、链上记录查询、比对判断和审计结果写入共同构成。随着审计数据规模从 100 条增加到 1000 条，总耗时同步上升，说明当前原型系统仍存在批量审计优化空间。批量审计耗时对比如图 5-6 所示，三组数据 5 轮均执行成功，表明系统在实验规模下能够稳定完成审计闭环。",
}


def is_heading(text: str) -> bool:
    marker = text.strip().split(maxsplit=1)[0] if text.strip() else ""
    return marker.count(".") >= 1 and all(part.isdigit() for part in marker.split("."))


def iter_paragraph_texts(doc: Document) -> list[str]:
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()]


def insert_after_element(element, parent, text: str, *, center: bool = False) -> Paragraph:
    new_p = OxmlElement("w:p")
    element.addnext(new_p)
    paragraph = Paragraph(new_p, parent)
    paragraph.text = text
    if center:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in paragraph.runs:
        run.font.name = "宋体"
        run.font.size = Pt(10.5)
    return paragraph


def find_section_end_paragraph(doc: Document, heading: str) -> Paragraph:
    start = None
    for index, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == heading:
            start = index
            break
    if start is None:
        raise ValueError(f"未找到章节标题：{heading}")

    end = start
    for index in range(start + 1, len(doc.paragraphs)):
        text = doc.paragraphs[index].text.strip()
        if is_heading(text):
            break
        if text:
            end = index
    return doc.paragraphs[end]


def find_table_by_header(doc: Document, header: list[str]) -> Table:
    for table in doc.tables:
        if not table.rows:
            continue
        cells = [cell.text.strip() for cell in table.rows[0].cells]
        if cells == header:
            return table
    raise ValueError(f"未找到表格：{header}")


def add_figure_placeholder_after(anchor_element, parent, figure: dict) -> None:
    placeholder = f"【待生成并插入图片：image/{figure['filename']}；图片内部不包含论文图题，图题仅保留在正文下方】"

    transition = insert_after_element(anchor_element, parent, figure["transition"])
    image_slot = insert_after_element(transition._p, parent, placeholder, center=True)
    insert_after_element(image_slot._p, parent, figure["caption"], center=True)


def figure_caption_exists(doc: Document, caption: str) -> bool:
    return any(p.text.strip() == caption for p in doc.paragraphs)


def add_placeholders(doc: Document) -> int:
    inserted = 0
    for figure in FIGURES:
        if figure_caption_exists(doc, figure["caption"]):
            continue
        if figure["anchor_type"] == "section_end":
            anchor = find_section_end_paragraph(doc, figure["section"])
            add_figure_placeholder_after(anchor._p, anchor._parent, figure)
        elif figure["anchor_type"] == "after_table":
            table = find_table_by_header(doc, figure["table_header"])
            add_figure_placeholder_after(table._tbl, table._parent, figure)
        else:
            raise ValueError(f"未知插入方式：{figure['anchor_type']}")
        inserted += 1
    return inserted


def rename_existing_captions_and_references(doc: Document) -> int:
    changed = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text in EXISTING_CAPTION_RENAMES:
            paragraph.text = EXISTING_CAPTION_RENAMES[text]
            changed += 1
            continue
        if text in TEXT_REPLACEMENTS:
            paragraph.text = TEXT_REPLACEMENTS[text]
            changed += 1
    return changed


def build_prompt_markdown() -> str:
    lines: list[str] = [
        "# 论文新增插图 AI 生成提示词（无标题版）",
        "",
        "论文题目：《基于区块链的可信任务日志审计系统设计与实现》",
        "",
        "本文件用于生成新增论文插图。所有图片请统一保存到 `D:\\aaaProject\\graduation-project\\image` 目录。图片本体内部不要出现“图 X-X ……”这类论文图题，图题只保留在 Word 正文图片下方。",
        "",
        "## 通用要求",
        "",
        "- 输出格式：PNG。",
        "- 推荐尺寸：1920x1080，横向 16:9。",
        "- 图片风格：本科软件工程毕业论文技术图，正式、清晰、克制、学术化。",
        "- 背景：白色或极浅灰色。",
        "- 文字：允许使用简短中文标签，但不要把论文图题写进图片内部。",
        "- 禁止内容：水印、二维码、真实公司 Logo、比特币 Logo、人物、卡通装饰、复杂深色背景、论文中不存在的技术组件。",
        "- 如果生成工具不能稳定生成中文，优先生成无文字结构图，再用 PPT、draw.io、ProcessOn 或 Word 手动添加中文标签。",
        "",
    ]

    for figure in FIGURES:
        lines.extend(
            [
                f"## {figure['caption']}",
                "",
                f"- 建议保存文件：`image/{figure['filename']}`",
                f"- 对应章节：`{figure['section']}`",
                "- 图片内部标题：不要出现论文图题，不要出现“图 X-X”。",
                "- 图片内容要点：",
            ]
        )
        for item in figure["content"]:
            lines.append(f"  - {item}")
        lines.extend(
            [
                "",
                "### 直接复制给图片生成工具的提示词",
                "",
                "```text",
                figure["prompt"],
                "```",
                "",
                "### 负面提示词",
                "",
                "```text",
                "不要在图片内部生成论文图题，不要出现水印、二维码、真实公司 Logo、人物、卡通装饰、复杂深色背景、错别字、乱码，不要加入论文没有实现的功能或技术栈。",
                "```",
                "",
            ]
        )

    lines.extend(
        [
            "## 生成后插入说明",
            "",
            "1. 按上述文件名生成 PNG 并保存到 `image` 目录。",
            "2. Word 中已经预留了对应图片位置，位置处有“待生成并插入图片”占位文字。",
            "3. 图片生成后，可将占位文字替换为对应 PNG，保留下方的图题。",
            "4. 图片本体不需要标题，因为 Word 下方已经有正式图题。",
            "",
        ]
    )
    return "\n".join(lines)


def main() -> None:
    if not DOCX_PATH.exists():
        raise FileNotFoundError(DOCX_PATH)

    backup = DOCX_PATH.with_name(
        f"{DOCX_PATH.stem}-before-additional-image-positions-{datetime.now().strftime('%Y%m%d-%H%M%S')}{DOCX_PATH.suffix}"
    )
    shutil.copy2(DOCX_PATH, backup)

    doc = Document(DOCX_PATH)
    renamed = rename_existing_captions_and_references(doc)
    inserted = add_placeholders(doc)
    doc.save(DOCX_PATH)

    PROMPT_PATH.write_text(build_prompt_markdown(), encoding="utf-8")

    REPORT_PATH.write_text(
        "# 论文新增图片位置修改报告\n\n"
        f"- 论文文件：`{DOCX_PATH}`\n"
        f"- 备份文件：`{backup}`\n"
        f"- 新增图片占位数量：{inserted}\n"
        f"- 已调整原有图题或图文引用数量：{renamed}\n"
        f"- 新增图片提示词文件：`{PROMPT_PATH}`\n\n"
        "## 新增图位\n\n"
        + "\n".join(f"- {item['caption']}：`image/{item['filename']}`，位置：{item['section']}" for item in FIGURES)
        + "\n\n## 说明\n\n"
        "- 图片内部不放论文图题，图题保留在 Word 正文中。\n"
        "- 当前 Word 先放置图片占位文字；生成 PNG 后再替换占位文字即可。\n",
        encoding="utf-8",
    )

    print(f"backup={backup}")
    print(f"inserted={inserted}")
    print(f"renamed={renamed}")
    print(f"prompt_file={PROMPT_PATH}")
    print(f"report={REPORT_PATH}")


if __name__ == "__main__":
    main()
