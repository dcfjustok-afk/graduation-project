from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt


DOCX_PATH = Path("output/Graduation-thesis.docx")
IMAGE_DIR = Path("image")


FIGURES = [
    {
        "heading": "3.1.1 链下日志存储与链上哈希存证模型",
        "caption": "图 3-1 链上链下混合存储模型",
        "image": "fig-3-1-chain-offchain-hybrid-storage-model.png",
    },
    {
        "heading": "3.1.2 日志采集、存证、审计、告警闭环流程",
        "caption": "图 3-2 可信日志审计闭环流程",
        "image": "fig-3-2-trusted-log-audit-loop-flow.png",
    },
    {
        "heading": "3.4.1 数据库哈希、链上哈希与重算哈希三方比对",
        "caption": "图 3-3 三方哈希比对流程",
        "image": "fig-3-3-three-party-hash-comparison-flow.png",
    },
    {
        "heading": "4.4 系统总体架构设计",
        "caption": "图 4-1 系统总体架构图",
        "image": "4.1系统总体架构图-去标题.png",
    },
    {
        "heading": "4.6 系统数据流与业务流程设计",
        "caption": "图 4-2 系统数据流图",
        "image": "图4-2-系统数据流图-去标题.png",
    },
    {
        "heading": "5.4.1 系统总览与日志中心实现",
        "caption": "图 5-1 系统前端总览页面",
        "image": "fig-5-1-system-frontend-overview.png",
    },
    {
        "heading": "5.4.2 审计管理与告警管理实现",
        "caption": "图 5-2 审计管理与告警管理页面",
        "image": "fig-5-2-audit-alert-management-page.png",
    },
    {
        "heading": "5.6.2 批量审计性能实验",
        "caption": "图 5-3 批量审计耗时对比图",
        "image": "fig-5-3-batch-audit-time-comparison.png",
    },
]


def is_heading(text: str) -> bool:
    text = text.strip()
    if not text:
        return False
    if text.startswith("第 ") and "章" in text:
        return True
    marker = text.split(maxsplit=1)[0]
    return marker.count(".") >= 1 and all(part.isdigit() for part in marker.split("."))


def paragraph_texts(doc: Document) -> list[str]:
    return [p.text.strip() for p in doc.paragraphs]


def remove_existing_figure_blocks(doc: Document) -> int:
    """Remove picture paragraphs and captions that this script previously inserted."""
    captions = {item["caption"] for item in FIGURES}
    removed = 0
    idx = 0
    while idx < len(doc.paragraphs):
        paragraph = doc.paragraphs[idx]
        text = paragraph.text.strip()
        if text in captions:
            # The picture paragraph is normally right before the caption.
            if idx > 0:
                prev = doc.paragraphs[idx - 1]
                if prev._element.xpath(".//w:drawing") or prev._element.xpath(".//w:pict"):
                    delete_paragraph(prev)
                    removed += 1
                    idx -= 1
            delete_paragraph(paragraph)
            removed += 1
            continue
        idx += 1
    return removed


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)
    paragraph._p = paragraph._element = None


def insert_after_paragraph(paragraph, text: str = ""):
    from docx.text.paragraph import Paragraph

    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.text = text
    return new_para


def find_section_end_index(doc: Document, heading: str) -> int:
    paragraphs = doc.paragraphs
    start = None
    for index, paragraph in enumerate(paragraphs):
        if paragraph.text.strip() == heading:
            start = index
            break
    if start is None:
        raise ValueError(f"未找到章节标题：{heading}")

    end = start
    for index in range(start + 1, len(paragraphs)):
        text = paragraphs[index].text.strip()
        if is_heading(text):
            break
        if text:
            end = index
    return end


def set_run_font(paragraph, size_pt: float = 10.5) -> None:
    for run in paragraph.runs:
        run.font.size = Pt(size_pt)
        run.font.name = "宋体"


def insert_figure(doc: Document, heading: str, image_path: Path, caption: str) -> None:
    end_index = find_section_end_index(doc, heading)
    anchor = doc.paragraphs[end_index]

    caption_para = insert_after_paragraph(anchor, caption)
    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(caption_para, 10.5)

    picture_para = insert_after_paragraph(anchor, "")
    picture_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = picture_para.add_run()
    run.add_picture(str(image_path), width=Inches(5.8))


def main() -> None:
    if not DOCX_PATH.exists():
        raise FileNotFoundError(DOCX_PATH)

    missing = [item["image"] for item in FIGURES if not (IMAGE_DIR / item["image"]).exists()]
    if missing:
        raise FileNotFoundError("缺少图片文件：" + "；".join(missing))

    backup = DOCX_PATH.with_name(
        f"{DOCX_PATH.stem}-before-image-insert-{datetime.now().strftime('%Y%m%d-%H%M%S')}{DOCX_PATH.suffix}"
    )
    shutil.copy2(DOCX_PATH, backup)

    doc = Document(DOCX_PATH)
    removed = remove_existing_figure_blocks(doc)

    inserted: list[str] = []
    for item in FIGURES:
        image_path = IMAGE_DIR / item["image"]
        insert_figure(doc, item["heading"], image_path, item["caption"])
        inserted.append(f"{item['caption']} -> {image_path.as_posix()}")

    doc.save(DOCX_PATH)

    report_path = Path("output/论文图片插入报告.md")
    report_path.write_text(
        "# 论文图片插入报告\n\n"
        f"- 论文文件：`{DOCX_PATH.as_posix()}`\n"
        f"- 备份文件：`{backup.as_posix()}`\n"
        f"- 本次插入图片数量：{len(inserted)}\n"
        f"- 清理旧图片/旧图题段落数量：{removed}\n\n"
        "## 插入明细\n\n"
        + "\n".join(f"- {line}" for line in inserted)
        + "\n",
        encoding="utf-8",
    )
    print(report_path.as_posix())
    print(f"backup={backup.as_posix()}")
    print(f"inserted={len(inserted)}")
    print(f"removed={removed}")


if __name__ == "__main__":
    main()
