from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "output" / "stage0-writing-materials.md"
TARGET = ROOT / "output" / "stage0-writing-materials.docx"


def set_font(run, size=10.5, bold=False):
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold


def add_paragraph(document, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.first_line_indent = Pt(21)
    run = paragraph.add_run(text)
    set_font(run)


def add_bullet(document, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.left_indent = Pt(18)
    paragraph.paragraph_format.first_line_indent = Pt(-18)
    run = paragraph.add_run("• " + text)
    set_font(run)


def add_numbered(document, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.left_indent = Pt(18)
    paragraph.paragraph_format.first_line_indent = Pt(-18)
    run = paragraph.add_run(text)
    set_font(run)


def add_heading(document, text, level):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.5
    if level == 1:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        size = 16
    elif level == 2:
        size = 15
    elif level == 3:
        size = 14
    else:
        size = 12
    run = paragraph.add_run(text)
    set_font(run, size=size, bold=True)


def add_table(document, rows):
    table = document.add_table(rows=1, cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for index, value in enumerate(rows[0]):
        cell = table.rows[0].cells[index]
        cell.text = value
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_font(run, bold=True)
    for row in rows[1:]:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
            for paragraph in cells[index].paragraphs:
                for run in paragraph.runs:
                    set_font(run)


def parse_table(lines, start_index):
    table_lines = []
    index = start_index
    while index < len(lines) and lines[index].strip().startswith("|"):
        table_lines.append(lines[index].strip())
        index += 1

    rows = []
    for line in table_lines:
        parts = [part.strip() for part in line.strip("|").split("|")]
        if all(part.replace("-", "").replace(":", "").strip() == "" for part in parts):
            continue
        rows.append(parts)
    return rows, index


def generate_docx():
    text = SOURCE.read_text(encoding="utf-8")
    document = Document()

    style = document.styles["Normal"]
    style.font.name = "宋体"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(10.5)

    lines = text.splitlines()
    index = 0
    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()
        if not stripped:
            index += 1
            continue

        if stripped.startswith("### "):
            add_heading(document, stripped[4:], 3)
        elif stripped.startswith("## "):
            add_heading(document, stripped[3:], 2)
        elif stripped.startswith("# "):
            add_heading(document, stripped[2:], 1)
        elif stripped.startswith("|"):
            rows, index = parse_table(lines, index)
            add_table(document, rows)
            continue
        elif stripped.startswith("- "):
            add_bullet(document, stripped[2:])
        elif stripped[:2].isdigit() and stripped[2:4] == ". ":
            add_numbered(document, stripped)
        elif stripped[:1].isdigit() and stripped[1:3] == ". ":
            add_numbered(document, stripped)
        else:
            add_paragraph(document, stripped)
        index += 1

    document.save(TARGET)


if __name__ == "__main__":
    generate_docx()
    print(TARGET)
