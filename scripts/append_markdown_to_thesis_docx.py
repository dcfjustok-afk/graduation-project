from pathlib import Path
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_TARGET = ROOT / "output" / "Graduation-thesis.docx"


def set_font(run, size=10.5, bold=False):
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold


def add_heading(document, text, level):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.5
    size = 16 if level == 1 else 15 if level == 2 else 12
    if level == 1:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    set_font(run, size=size, bold=True)


def add_paragraph(document, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Pt(21)
    paragraph.paragraph_format.line_spacing = 1.5
    run = paragraph.add_run(text)
    set_font(run)


def add_bullet(document, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Pt(18)
    paragraph.paragraph_format.first_line_indent = Pt(-18)
    paragraph.paragraph_format.line_spacing = 1.5
    run = paragraph.add_run("• " + text)
    set_font(run)


def add_numbered(document, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Pt(18)
    paragraph.paragraph_format.first_line_indent = Pt(-18)
    paragraph.paragraph_format.line_spacing = 1.5
    run = paragraph.add_run(text)
    set_font(run)


def add_center_placeholder(document, text):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.line_spacing = 1.5
    run = paragraph.add_run(text)
    set_font(run, bold=True)


def add_contract_method_table(document):
    rows = [
        ("合约方法", "主要功能", "审计阶段作用"),
        ("storeLog(taskId, logHash)", "写入日志哈希存证记录，并触发 LogStored 事件。", "形成链上存证依据。"),
        ("getLog(recordId)", "根据记录编号查询单条链上记录。", "支持按记录编号定位存证数据。"),
        ("getRecordIdsByTaskId(taskId)", "查询某一任务下全部链上记录编号。", "支持任务维度记录定位。"),
        ("getLogsByTaskId(taskId)", "查询某一任务下完整链上日志记录。", "用于审计阶段读取链上哈希。"),
        ("getTaskLogCount(taskId)", "统计指定任务下链上日志数量。", "辅助判断任务存证规模。"),
        ("getLogCount()", "统计合约中全部存证记录数量。", "辅助整体存证统计。"),
    ]
    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, value in enumerate(rows[0]):
        table.rows[0].cells[idx].text = value
    for row in rows[1:]:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_font(run, bold=(row == table.rows[0]))


def add_audit_status_table(document):
    rows = [
        ("审计状态", "判定条件", "含义"),
        ("passed", "expectedHash、actualHash 与 onChainHash 三方一致。", "日志内容与数据库记录、链上存证相互匹配，审计通过。"),
        ("failed", "actualHash 与 expectedHash 不一致，或链上记录与数据库记录不匹配。", "日志存在篡改风险或存证记录不一致，需要生成告警并复核。"),
        ("pending", "本地哈希与数据库记录一致，但链上记录缺失、不可用或暂未匹配。", "当前证据不完整，审计结论待进一步确认。"),
    ]
    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, value in enumerate(rows[0]):
        table.rows[0].cells[idx].text = value
    for row in rows[1:]:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_font(run, bold=(row == table.rows[0]))


def add_database_core_table(document):
    rows = [
        ("表名", "主要用途", "关键字段"),
        ("logs", "保存日志原文与采集来源信息。", "id、task_id、source_type、source_path、log_content、log_level、collected_at、status"),
        ("log_hash_records", "保存日志哈希及链上写入信息。", "log_id、task_id、log_hash、contract_address、transaction_hash、block_number、on_chain_status"),
        ("audit_records", "保存审计执行结果。", "log_id、log_hash_record_id、audit_status、expected_hash、actual_hash、audit_message、audited_at"),
        ("alerts", "保存异常告警记录。", "alert_type、severity、related_log_id、related_audit_id、title、description、status"),
        ("agent_states", "保存 Agent 运行状态。", "agent_name、source_path、last_offset、last_heartbeat_at、last_sync_at、status、error_message"),
    ]
    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, value in enumerate(rows[0]):
        table.rows[0].cells[idx].text = value
    for row in rows[1:]:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_font(run, bold=(row == table.rows[0]))


def add_project_structure_table(document):
    rows = [
        ("目录", "主要内容", "作用"),
        ("apps/web", "React、TypeScript、Ant Design 前端页面。", "展示总览、日志、审计和告警信息。"),
        ("apps/server", "Express 后端服务、业务逻辑、数据库访问和链上交互。", "负责日志接收、哈希计算、上链、审计和告警。"),
        ("apps/agent", "日志采集 Agent、文件读取、重试队列和状态同步。", "负责自动采集本地日志并提交后端。"),
        ("packages/contracts", "Hardhat 工程和 LogRegistry Solidity 合约。", "提供链上日志哈希存证能力。"),
        ("packages/shared", "共享类型和公共协议。", "统一前后端及 Agent 的数据结构。"),
        ("tests/performance", "性能测试脚本。", "支撑日志提交和批量审计实验。"),
        ("scripts", "仓库级验证脚本。", "支撑合约、后端、Agent 和前端的脚本化验证。"),
    ]
    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, value in enumerate(rows[0]):
        table.rows[0].cells[idx].text = value
    for row in rows[1:]:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_font(run, bold=(row == table.rows[0]))


def append_markdown(source, target):
    document = Document(target) if target.exists() else Document()
    style = document.styles["Normal"]
    style.font.name = "宋体"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(10.5)

    text = source.read_text(encoding="utf-8")
    if any(paragraph.text.strip() for paragraph in document.paragraphs) or document.tables:
        document.add_page_break()

    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("## "):
            add_heading(document, line[3:], 2)
        elif line.startswith("# "):
            add_heading(document, line[2:], 1)
        elif line.startswith("- "):
            add_bullet(document, line[2:])
        elif line == "表 3-2 LogRegistry 合约方法说明":
            add_center_placeholder(document, line)
            add_contract_method_table(document)
        elif line == "表 3-3 审计状态判定规则":
            add_center_placeholder(document, line)
            add_audit_status_table(document)
        elif line == "表 4-1 数据库核心表结构说明":
            add_center_placeholder(document, line)
            add_database_core_table(document)
        elif line == "表 5-1 系统工程目录说明":
            add_center_placeholder(document, line)
            add_project_structure_table(document)
        elif line.startswith("图 ") or line.startswith("表 "):
            add_center_placeholder(document, line)
        elif line[:2].isdigit() and line[2:4] == ". ":
            add_numbered(document, line)
        elif line[:1].isdigit() and line[1:3] == ". ":
            add_numbered(document, line)
        else:
            add_paragraph(document, line)

    document.save(target)


if __name__ == "__main__":
    if len(sys.argv) not in (2, 3):
        raise SystemExit("usage: python scripts/append_markdown_to_thesis_docx.py <markdown-path> [target-docx]")
    source_path = Path(sys.argv[1])
    if not source_path.is_absolute():
        source_path = ROOT / source_path
    target_path = Path(sys.argv[2]) if len(sys.argv) == 3 else DEFAULT_TARGET
    if not target_path.is_absolute():
        target_path = ROOT / target_path
    append_markdown(source_path, target_path)
    print(target_path)
