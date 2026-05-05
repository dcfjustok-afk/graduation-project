from __future__ import annotations

from pathlib import Path

from docx import Document


DOCX_PATH = Path("output/Graduation-thesis.docx")


REPLACEMENTS = {
    "1.3 本文主要研究内容": [
        "本文围绕基于区块链的可信任务日志审计系统展开研究，面向传统中心化日志易被篡改、删除或伪造且审计依据不足的问题，设计并实现包含日志采集、链下存储、链上存证、审计比对、异常告警和前端展示的原型系统。主要研究内容包括以下方面。",
        "第一，设计链上链下混合存储模型。系统将日志原文保存在 SQLite 中，将 SHA-256 哈希摘要和任务标识写入 LogRegistry 合约，以兼顾隐私保护、存储成本和审计可信性。第二，设计 LogRegistry 智能合约，保存 taskId、logHash、createdAt 和 submitter，并通过 LOGGER_ROLE 控制写入权限，提供任务维度查询能力。",
        "第三，设计三方哈希比对审计机制。系统在审计阶段重新计算 actualHash，并与数据库 expectedHash、链上 onChainHash 进行比对，形成 passed、failed 和 pending 状态。第四，实现 Agent 自动采集、后端服务、前端展示和告警联动，并通过合约测试、批量提交、批量审计和篡改检测实验验证系统闭环可运行。",
    ],
    "2.2 智能合约与访问控制技术": [
        "智能合约是部署在区块链上的程序化规则，能够在满足条件时自动执行，并将调用记录保存在链上。与普通后端程序相比，智能合约的执行过程和状态变化更容易被追溯，适合承担证据登记、权限校验和链上查询等任务。但合约部署后修改成本较高，因此其数据结构和权限边界需要在设计阶段明确。",
        "访问控制是智能合约安全设计的重要内容。在日志存证场景中，若任意地址都能写入日志哈希，攻击者可能提交伪造摘要，削弱链上证据的可信性。OpenZeppelin AccessControl 提供基于角色的授权机制，可通过 DEFAULT_ADMIN_ROLE 管理角色，通过 LOGGER_ROLE 限定日志写入主体。相比自定义布尔字段，该机制接口清晰、可复用性较好，适合 LogRegistry 合约的写入权限控制。",
    ],
    "2.5 本章小结": [
        "本章围绕系统涉及的关键技术进行了说明，主要包括区块链存证、智能合约与访问控制、哈希完整性校验，以及 React、Express、SQLite、Hardhat 等开发技术。这些内容为后续链上链下混合存储、LogRegistry 合约设计和三方哈希审计机制提供了技术基础。",
    ],
    "3.2.3 链上存证字段与链下记录映射关系": [
        "本系统采用链下业务数据与链上摘要数据相互映射的方式组织日志证据。logs 表保存日志原文及采集来源，是审计时重新计算 actualHash 的基础；log_hash_records 表保存提交阶段生成的 log_hash、contract_address、transaction_hash、block_number 和 on_chain_status 等信息；LogRegistry 合约保存 taskId、logHash、createdAt 和 submitter，形成链上存证依据。",
        "其中，logs.id 与 log_hash_records.log_id 对应，用于将日志原文和哈希记录关联起来；task_id 对应链上 taskId，用于按任务归集日志；log_hash 对应链上 logHash，用于后续完整性比对；contract_address 记录写入时使用的合约地址，支持历史合约地址回溯；transaction_hash 和 block_number 用于定位链上交易和区块；on_chain_status 表示上链结果。",
        "通过上述映射关系，系统能够在审计阶段同时取得日志原文、数据库哈希和链上哈希。若 expectedHash、actualHash 和 onChainHash 一致，则说明链上链下证据匹配；若不一致，则可进一步结合交易哈希、区块号和合约地址定位异常来源。表 3-1 链上链下字段映射关系",
    ],
    "3.3.1 合约数据结构设计": [
        "LogRegistry 合约以 LogRecord 结构体保存链上日志存证信息，主要字段包括 taskId、logHash、createdAt 和 submitter。taskId 表示任务标识，用于将同一任务产生的多条日志归集；logHash 表示日志内容经 SHA-256 计算后的摘要，是后续完整性校验的核心字段；createdAt 记录链上写入时间；submitter 记录提交该日志哈希的账户地址。",
        "合约内部通过 records 数组保存全部 LogRecord，并以数组下标或记录编号支持单条记录查询。为提高任务维度查询效率，合约还维护 taskIdToRecordIds 映射，将同一 taskId 对应的记录编号集中保存。审计时，系统既可以通过记录编号查询单条链上存证，也可以根据 taskId 获取某一任务下的全部链上记录。",
        "该数据结构没有保存日志原文，只保存必要摘要和元数据，能够降低链上存储压力，也避免敏感日志内容直接暴露。records 数组保证全局记录可追溯，taskIdToRecordIds 映射支持任务维度审计，两者共同满足日志编号查询和任务回溯查询需求。",
    ],
    "3.3.2 基于 AccessControl 的写入权限控制": [
        "日志存证合约不能允许任意地址写入。如果链上写入缺少权限限制，攻击者或无关账户可能提交伪造 taskId 和 logHash，导致链上记录与真实日志无关，从而影响审计依据的可信性。因此，LogRegistry 合约需要在写入入口设置安全边界，只允许受信任的日志提交主体调用 storeLog。",
        "本系统基于 OpenZeppelin AccessControl 实现角色控制。DEFAULT_ADMIN_ROLE 用于管理角色授权，LOGGER_ROLE 用于控制日志哈希写入权限。后端服务使用具备 LOGGER_ROLE 的账户调用 storeLog，将日志哈希写入链上；未获得该角色的账户即使知道合约地址和方法，也不能完成有效写入。",
        "与自定义权限字段相比，AccessControl 提供成熟的角色管理接口和事件记录，便于授权、撤销和审计。该设计将链下业务系统的提交行为与链上权限边界对应起来，使 LogRegistry 合约不仅保存证据摘要，也能够限制可信写入来源。",
    ],
    "3.4.1 数据库哈希、链上哈希与重算哈希三方比对": [
        "三方哈希比对是本系统审计机制的核心。系统在日志提交阶段计算日志内容的 SHA-256 摘要，并将该值写入 log_hash_records，作为 expectedHash；同时将对应哈希写入 LogRegistry 合约，形成链上 onChainHash。审计阶段，系统再次读取 logs 表中的日志原文并重新计算 actualHash。",
        "审计逻辑以 expectedHash、actualHash 和 onChainHash 的一致性为判断依据。若三者完全一致，说明当前日志原文、数据库存证记录和链上存证记录相互匹配，审计结果为 passed。若 actualHash 与 expectedHash 不一致，说明链下日志原文可能被修改；若 expectedHash 与 onChainHash 不一致，则说明数据库记录与链上存证存在差异，审计结果为 failed。",
        "当链上记录缺失、合约地址不可用或链上查询结果不完整时，系统不直接判定日志被篡改，而是将结果置为 pending，并保留审计说明。该设计能够区分内容篡改、证据缺失和链上环境异常，避免单一数据库校验带来的可信性不足。图 3-3 三方哈希比对流程",
    ],
    "3.5.1 审计状态判定规则": [
        "系统将审计结果划分为 passed、failed 和 pending 三种状态。passed 表示 expectedHash、actualHash 与 onChainHash 完全一致，说明日志原文、数据库哈希和链上哈希能够相互印证。该状态下，系统将审计记录写入 audit_records，并保留审计时间和说明信息。",
        "failed 表示审计过程中发现明确不一致。典型情况包括 actualHash 与 expectedHash 不一致，说明 logs 表中的日志内容可能被修改；或者 expectedHash 与 onChainHash 不一致，说明链下存证记录与链上记录不匹配。此时系统会生成 failed 审计记录，并触发 hash_mismatch 告警。",
        "pending 表示当前证据不足以得出最终结论，常见于链上记录缺失、合约地址不可用或合约代码不存在等情况。pending 不等同于审计通过，而是提示需要补充链上证据或检查部署环境。表 3-2 审计状态判定规则",
    ],
    "3.6 本章小结": [
        "本章围绕可信日志存证与审计机制展开设计，说明了链下日志原文与链上哈希摘要结合的混合存储模型，设计了日志标准化封装、SHA-256 摘要生成和链上链下字段映射关系，并围绕 LogRegistry 合约阐述了数据结构、LOGGER_ROLE 权限控制和任务维度查询方法。",
        "在审计机制方面，本章重点说明了 expectedHash、actualHash 与 onChainHash 的三方哈希比对，给出了历史合约地址回溯、provider.getCode(address) 合约代码存在性校验、审计状态判定和 hash_mismatch 告警流程，为后续系统总体设计与工程实现奠定基础。",
    ],
    "4.4 系统总体架构设计": [
        "本系统采用分层和模块化架构，由 Agent、Server、SQLite、LogRegistry 和 Web 五类核心部分组成，分别承担日志采集、业务处理、链下存储、链上存证和前端展示职责。Agent 位于数据流入口，负责从本地日志文件读取新增内容，并封装任务标识、来源路径、日志级别和采集时间等字段后提交 Server。",
        "Server 是系统业务处理中心，负责接收日志、写入 SQLite、计算 SHA-256 哈希、调用 LogRegistry 合约以及执行审计任务。SQLite 保存 logs、log_hash_records、audit_records、alerts 和 agent_states 等数据，为查询和展示提供基础。LogRegistry 仅保存日志哈希和必要元数据，通过 LOGGER_ROLE 控制写入权限。",
        "Web 前端面向用户展示系统总览、日志列表、审计结果和告警信息。从数据流看，日志先由 Agent 提交至 Server，再写入 SQLite 并完成链上存证；审计时，Server 同时读取链下原文、数据库哈希和链上哈希，完成三方比对并将结果展示到前端。图 4-1 系统总体架构图",
    ],
    "4.6 系统数据流与业务流程设计": [
        "系统数据流主要包括日志采集与存证流程、审计与告警流程。日志采集与存证流程从 Agent 读取日志文件开始，Agent 根据本地偏移量只采集新增内容，并将 taskId、sourcePath、logContent、logLevel 和 collectedAt 等字段提交给 Server。Server 接收后将日志写入 logs 表，随后对 log_content 计算 SHA-256 哈希。",
        "哈希生成后，Server 调用 LogRegistry 的 storeLog(taskId, logHash) 方法完成链上写入，并将 log_id、task_id、log_hash、contract_address、transaction_hash、block_number 和 on_chain_status 写入 log_hash_records。至此，日志原文保存在链下，哈希摘要保存在链上，形成可审计的存证记录。",
        "审计与告警流程由审计任务触发。Server 读取 logs 表中的日志原文并重新计算 actualHash，读取 log_hash_records 中的 expectedHash，再根据合约地址和任务标识查询链上 onChainHash。若三者一致，结果为 passed；若哈希不一致，结果为 failed，并生成 hash_mismatch 告警；若链上记录不可用，则进入 pending。最终，Web 前端展示日志、审计和告警结果。图 4-2 系统数据流图",
    ],
    "5.1 系统开发环境与工程结构": [
        "系统采用 Monorepo 工程结构组织代码，将前端、后端、Agent、智能合约、共享类型和测试脚本统一放在同一仓库下管理。主要目录包括 apps/web、apps/server、apps/agent、packages/contracts、packages/shared 和 tests/performance 等。该结构便于统一依赖、复用类型定义和维护跨模块接口。",
        "apps/web 基于 React、TypeScript 和 Ant Design 实现前端页面；apps/server 基于 Node.js 和 Express 实现日志接收、数据库访问、哈希计算、合约调用、审计和告警逻辑；apps/agent 负责日志文件增量采集、偏移量保存和失败重试；packages/contracts 使用 Hardhat、Solidity、OpenZeppelin AccessControl 和 Ethers 实现 LogRegistry 合约。",
        "packages/shared 用于存放共享类型和协议，tests/performance 用于支撑日志批量提交和批量审计实验。Monorepo 结构使系统模块边界清晰，有利于毕业设计中的工程说明、功能验证和后续扩展。表 5-1 工程目录说明",
    ],
    "5.2.1 增量读取与偏移量持久化实现": [
        "日志采集 Agent 的核心任务是从指定日志文件中持续读取新增内容，并避免重复提交。系统通过 fileReader 读取文件内容，通过 offsetStore 保存每个 sourcePath 的 last_offset。当 Agent 再次启动或继续采集时，能够从上次记录的偏移量之后读取新增日志，从而实现增量采集。",
        "logCollector 负责将原始日志行封装为统一数据结构，包括 taskId、sourceType、sourcePath、logContent、logLevel 和 collectedAt 等字段。封装后的日志提交给后端接口，由后端完成入库、哈希计算和上链。偏移量只有在日志提交成功后更新，避免网络异常导致未提交内容被跳过。",
        "该实现使 Agent 能够适应日志文件持续增长的场景，也为审计系统提供稳定的数据入口。通过偏移量持久化，系统在重启后仍能继续采集新增日志，降低重复采集和遗漏采集风险。",
    ],
    "5.2.2 失败重试与状态同步实现": [
        "日志采集过程中可能出现后端暂时不可用、网络连接失败或接口返回异常等情况。为降低临时错误对采集链路的影响，Agent 使用 retryQueue 保存提交失败的日志数据，并在后续周期中重新尝试发送。只有当后端确认接收成功后，相关日志才会从重试队列中移除。",
        "logAgent 还会同步 Agent 运行状态，包括采集来源、当前偏移量、最后心跳时间、同步时间、运行状态和错误信息等。这些状态写入 agent_states 表后，前端或后端可以查看 Agent 是否正常工作，也便于定位采集异常。",
        "失败重试与状态同步增强了采集模块的可用性。即使短时间内后端不可访问，系统也能保留待提交日志，并在服务恢复后继续提交，从而保证采集、存证和审计闭环的连续性。",
    ],
    "5.3.3 审计记录与告警数据表实现": [
        "审计执行结果主要写入 audit_records 表。该表记录 log_id、log_hash_record_id、audit_status、expected_hash、actual_hash、audit_message 和 audited_at 等字段，用于保存每次审计的状态和依据。通过保留 expectedHash 与 actualHash，系统能够说明审计判断来源，而不是只保存最终结果。",
        "当审计状态为 failed 时，后端服务会生成 hash_mismatch 类型告警并写入 alerts 表。告警记录包含 related_log_id、related_audit_id、标题、描述、严重程度和处理状态等字段，用于将异常日志与审计记录关联起来。这样，用户在前端查看告警时，可以追溯到具体日志和审计过程。",
        "audit_records 与 alerts 的组合使系统形成从检测到提示的闭环。审计记录保存技术判断，告警记录面向用户提示风险，两者共同支撑日志篡改检测后的追踪和展示。",
    ],
    "5.6.3 篡改检测实验与结果分析": [
        "篡改检测实验用于验证系统能否识别日志内容被修改后的异常情况。实验先通过正常流程提交日志，使日志原文写入 logs 表，日志哈希写入 log_hash_records 和 LogRegistry；随后修改日志内容并触发审计。实验结果显示，auditStatus 为 failed，alertGenerated 为 true。",
        "该结果说明，当日志内容被修改后，审计阶段重新计算得到的 actualHash 与数据库 expectedHash 或链上 onChainHash 不一致，系统能够判定审计失败并生成告警。该实验验证了链上链下混合存储、三方哈希比对和 hash_mismatch 告警机制的有效性。",
    ],
    "6.1 研究工作总结": [
        "本文围绕可信任务日志审计系统完成了从机制设计到工程实现的工作。针对中心化日志易被篡改、删除或伪造的问题，系统采用链下保存日志原文、链上保存哈希摘要的方式，将日志内容与区块链存证结合起来，形成可核验的审计依据。",
        "在系统实现方面，本文设计了 Agent 增量采集、Express 后端服务、SQLite 链下存储、LogRegistry 智能合约、审计告警和 Web 可视化模块。审计阶段通过 expectedHash、actualHash 与 onChainHash 三方比对判断日志状态，并在 failed 时生成 hash_mismatch 告警。",
        "在验证方面，系统完成了合约测试、接口回归、Agent 采集、前端构建、批量提交、批量审计和篡改检测实验。实验结果表明，系统能够运行完整闭环，并能识别日志内容被修改的情况。",
    ],
    "6.2 系统创新点": [
        "本文的创新点主要体现在工程机制组合与可信审计闭环设计上。第一，系统采用链下日志原文与链上哈希摘要结合的轻量存证模型，在避免日志原文直接上链的同时，保留链上可验证证据，兼顾隐私、成本和审计需求。",
        "第二，系统设计 expectedHash、actualHash 和 onChainHash 三方哈希比对机制，将数据库记录、当前日志内容和链上存证同时纳入审计判断，增强了日志完整性校验的可信依据。第三，系统在工程实践中加入历史合约地址回溯和 provider.getCode(address) 合约代码存在性校验，减少本地链重启和多次部署导致的误判。",
        "第四，系统实现了采集、存证、审计、告警和展示闭环。该创新不属于理论突破，而是在毕业设计原型中将区块链存证、日志采集和异常告警整合为可运行系统。",
    ],
    "6.3 存在不足": [
        "当前系统仍存在一定局限。首先，实验环境主要基于本地 Hardhat 链，尚未接入公开测试网或联盟链，因此链上共识环境、节点稳定性和实际部署成本仍未充分验证。其次，SQLite 更适合原型系统和小规模实验，在高并发日志写入和大规模数据分析场景下能力有限。",
        "此外，批量审计流程以顺序处理为主，随着数据规模增加，审计耗时明显上升，后续仍需引入异步队列、批处理或并发审计优化。权限体系目前主要围绕 LOGGER_ROLE 展开，对多角色用户、审计人员权限和告警处置流程支持不足。系统也尚未深入分析日志语义，异常判断主要依赖哈希一致性。",
    ],
    "6.4 后续优化方向": [
        "后续可从部署环境、性能和安全治理三个方面优化系统。首先，可将 LogRegistry 部署到联盟链或公开测试网，验证多节点环境下的存证稳定性和交易确认情况。其次，可引入更完整的多角色权限体系，区分日志提交者、审计人员、管理员和告警处理人员，使系统更接近实际业务场景。",
        "在性能方面，可设计异步审计队列、批量哈希计算和并发链上查询机制，降低大规模审计耗时。数据存储层也可从 SQLite 扩展到更适合生产环境的数据库。在功能方面，可结合日志语义分析和风险分级，对错误日志、异常任务和高风险操作进行更细粒度识别，并通过前端提供更完善的统计图表和告警追踪能力。",
    ],
    "6.5 本章小结": [
        "本章总结了系统研究工作、主要创新点、存在不足和后续优化方向。总体来看，本文实现了基于区块链的可信任务日志审计原型系统，验证了链上链下混合存储、三方哈希比对和异常告警机制的可行性。同时，系统仍需在真实链环境、性能优化、权限体系和日志语义分析方面继续完善。",
    ],
}


def is_heading(text: str) -> bool:
    text = text.strip()
    if text in {"摘要", "Abstract", "致谢", "参考文献"}:
        return True
    if text.startswith("关键词：") or text.startswith("Keywords:"):
        return True
    if text.startswith("第 ") and "章" in text:
        return True
    parts = text.split(maxsplit=1)
    if not parts:
        return False
    marker = parts[0]
    return marker.count(".") >= 1 and all(part.isdigit() for part in marker.split("."))


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)
    paragraph._p = paragraph._element = None


def non_empty_paragraphs(doc: Document):
    return [(idx, p) for idx, p in enumerate(doc.paragraphs) if p.text.strip()]


def set_after_heading(doc: Document, heading: str, replacement: list[str]) -> None:
    entries = non_empty_paragraphs(doc)
    start_pos = None
    for pos, (_idx, paragraph) in enumerate(entries):
        if paragraph.text.strip() == heading:
            start_pos = pos
            break
    if start_pos is None:
        raise ValueError(f"heading not found: {heading}")

    body_positions: list[int] = []
    for pos in range(start_pos + 1, len(entries)):
        text = entries[pos][1].text.strip()
        if is_heading(text):
            break
        body_positions.append(pos)

    if not body_positions:
        raise ValueError(f"section has no body: {heading}")

    for offset, pos in enumerate(body_positions):
        paragraph = entries[pos][1]
        if offset < len(replacement):
            paragraph.text = replacement[offset]
        else:
            remove_paragraph(paragraph)


def table_chars(doc: Document) -> int:
    total = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                total += len(cell.text.strip())
    return total


def main() -> None:
    doc = Document(DOCX_PATH)
    for heading, replacement in REPLACEMENTS.items():
        set_after_heading(doc, heading, replacement)
    doc.save(DOCX_PATH)

    final_doc = Document(DOCX_PATH)
    paras = [p.text.strip() for p in final_doc.paragraphs if p.text.strip()]
    full = "\n".join(paras)
    report = Path("output/stage8-30k-final-tighten-report.md")
    report.write_text(
        "# 3 万字左右最终压缩报告\n\n"
        f"- 段落字符数：{sum(len(p) for p in paras)}\n"
        f"- 表格字符数：{table_chars(final_doc)}\n"
        f"- 合计估算字符数：{sum(len(p) for p in paras) + table_chars(final_doc)}\n"
        f"- 段落数：{len(paras)}\n"
        f"- 表格数：{len(final_doc.tables)}\n"
        f"- 问号数量：{full.count('?')}\n"
        f"- 本节关键信息摘要数量：{full.count('本节关键信息摘要')}\n",
        encoding="utf-8",
    )
    print(report.as_posix())


if __name__ == "__main__":
    main()
