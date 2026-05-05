from __future__ import annotations

from pathlib import Path

from docx import Document


DOCX_PATH = Path("output/Graduation-thesis.docx")


CHAPTERS = [
    "第 1 章 绪论",
    "第 2 章 相关理论与关键技术",
    "第 3 章 基于区块链的可信日志存证与审计机制设计",
    "第 4 章 系统需求分析与总体设计",
    "第 5 章 系统实现与实验验证",
    "第 6 章 总结与展望",
]


FIGURES = [
    "图 3-1 链上链下混合存储模型",
    "图 3-2 可信日志审计闭环流程",
    "图 3-3 三方哈希比对流程",
    "图 4-1 系统总体架构图",
    "图 4-2 系统数据流图",
    "图 5-1 系统前端总览页面",
    "图 5-2 审计管理与告警管理页面",
    "图 5-3 批量审计耗时对比图",
]


KEYS = [
    "LogRegistry",
    "LOGGER_ROLE",
    "provider.getCode",
    "contract_address",
    "expectedHash",
    "actualHash",
    "onChainHash",
    "hash_mismatch",
    "107.03",
    "9.33",
    "3067.77",
    "15659.44",
    "35032.13",
    "auditStatus",
    "alertGenerated",
]


def main() -> None:
    doc = Document(DOCX_PATH)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    full = "\n".join(paragraphs)

    print(f"file={DOCX_PATH.resolve()}")
    print(f"paragraphs={len(paragraphs)}")
    print(f"chars={sum(len(p) for p in paragraphs)}")
    print(f"inline_shapes={len(doc.inline_shapes)}")
    print(f"tables={len(doc.tables)}")
    print(f"replacement_char={full.count(chr(0xfffd))}")
    print(f"question_count={full.count('?')}")
    print(f"section_summary_count={full.count('本节关键信息摘要')}")
    print(f"first_person_我={full.count('我')}")
    print(f"first_person_我们={full.count('我们')}")
    print()

    print("CHAPTERS")
    for chapter in CHAPTERS:
        print(f"{chapter}\t{full.count(chapter)}")
    print()

    print("FIGURES")
    for figure in FIGURES:
        print(f"{figure}\t{full.count(figure)}")
    print()

    print("KEYS")
    for key in KEYS:
        print(f"{key}\t{full.count(key)}")
    print()

    print("TABLE_OR_FIGURE_LINES")
    for index, text in enumerate(paragraphs, start=1):
        if text.startswith(("图 ", "表 ")):
            print(f"{index:03d}: {text}")


if __name__ == "__main__":
    main()
