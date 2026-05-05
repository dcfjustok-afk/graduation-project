from __future__ import annotations

import re
import shutil
from collections import Counter
from datetime import datetime
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "output" / "Graduation-thesis.docx"
CLEAN_SOURCE_PATH = ROOT / "output" / "Graduation-thesis-before-reference-insert-20260506-000358.docx"
REPORT_PATH = ROOT / "output" / "论文参考文献插入校验报告.md"


REPLACEMENTS = {
    "随着政务服务、企业管理、自动化运维和任务调度系统的广泛应用，系统运行过程中会持续产生大量日志数据。日志记录了任务执行过程、接口调用结果、异常错误、用户操作和状态变化，是安全审计、故障排查、责任追溯和运行分析的重要依据。当系统发生异常或安全事件时，日志通常是还原执行过程、定位问题来源和判断责任边界的基础材料。因此，日志数据是否真实、完整、可追溯，直接影响后续审计结论的可信程度。":
    "随着政务服务、企业管理、自动化运维和任务调度系统的广泛应用，系统运行过程中会持续产生大量日志数据。日志记录了任务执行过程、接口调用结果、异常错误、用户操作和状态变化，是安全审计、故障排查、责任追溯和运行分析的重要依据[1]。当系统发生异常或安全事件时，日志通常是还原执行过程、定位问题来源和判断责任边界的基础材料。因此，日志数据是否真实、完整、可追溯，直接影响后续审计结论的可信程度。",

    "传统日志管理方式多采用中心化存储模式，日志原文通常保存在服务器文件系统、数据库或集中日志平台中。该方式部署简单、查询方便，也便于与业务系统集成，但在可信审计场景下存在明显不足：日志原文、哈希记录和审计结果往往处于同一管理域内，一旦管理权限被滥用或数据库被攻击，日志内容可能被修改、删除或伪造，事后难以提供独立证据说明日志是否保持原始状态。基于这一问题，本文围绕“基于区块链的可信任务日志审计系统设计与实现”展开研究，构建从日志采集、链下存储、链上存证、审计比对、异常告警到前端展示的原型系统。":
    "传统日志管理方式多采用中心化存储模式，日志原文通常保存在服务器文件系统、数据库或集中日志平台中。该方式部署简单、查询方便，也便于与业务系统集成，但在可信审计场景下存在明显不足：日志原文、哈希记录和审计结果往往处于同一管理域内，一旦管理权限被滥用或数据库被攻击，日志内容可能被修改、删除或伪造，事后难以提供独立证据说明日志是否保持原始状态。已有研究将区块链用于可信日志存储与验证，通过链上数据指纹与本地数据指纹比对识别日志内容被修改的情况[2]。区块链通过密码学结构和分布式账本机制，为可信数据记录提供了新的技术路径[3]。基于这一问题，本文围绕“基于区块链的可信任务日志审计系统设计与实现”展开研究，构建从日志采集、链下存储、链上存证、审计比对、异常告警到前端展示的原型系统。",

    "近年来，数据可信存证、日志审计和智能合约可信管理等方向持续发展。区块链存证研究强调链上记录的可追溯性，日志审计研究关注日志采集、集中管理和完整性校验，智能合约研究则为链上规则固化和权限控制提供了实现路径。相关研究普遍认为，大规模原文数据并不适合直接上链，而更适合通过哈希摘要、业务标识和时间信息进行轻量化存证。":
    "近年来，数据可信存证、日志审计和智能合约可信管理等方向持续发展。从可信数据管理角度看，区块链可在不完全可信环境中增强数据完整性、可追溯性和不可抵赖性[4]。区块链在信息安全领域的研究已覆盖认证、访问控制和数据保护等方向[5]。区块链存证研究强调链上记录的可追溯性，日志审计研究关注日志采集、集中管理和完整性校验，智能合约研究则为链上规则固化和权限控制提供了实现路径。相关研究普遍认为，大规模原文数据并不适合直接上链，而更适合通过哈希摘要、业务标识和时间信息进行轻量化存证。",

    "区块链存证技术是区块链在数据可信证明领域的重要应用方向。相关研究表明，区块链由于具备账本共享、记录可追溯和历史数据难以篡改等特点，可以用于保存电子证据、数据摘要、交易凭证和关键操作记录。在数据完整性保护场景中，区块链常被用作外部可信记录载体，为链下数据提供时间顺序和摘要校验依据。从存证方式看，常见方案主要包括完整数据上链和摘要数据上链两类。完整数据上链能够使数据直接成为链上记录，便于后续查询和验证，但在日志、文件、业务数据等规模较大的场景中，会带来较高的存储成本和处理开销，也可能导致敏感信息暴露。":
    "区块链存证技术是区块链在数据可信证明领域的重要应用方向。区块链数字取证研究关注证据获取、保全和呈现过程中的可信记录与链下扩展存储[6]。记录可信并不只由区块链技术本身决定，还依赖证据生成、登记、保存和验证流程之间的配合[7]。相关研究表明，区块链由于具备账本共享、记录可追溯和历史数据难以篡改等特点，可以用于保存电子证据、数据摘要、交易凭证和关键操作记录。在数据完整性保护场景中，区块链常被用作外部可信记录载体，为链下数据提供时间顺序和摘要校验依据；区块链也可通过区块结构、交易记录和共识机制为链下数据提供可追溯的外部参照[8]。在数据溯源研究中，区块链被用于保存链下数据的关键状态和追溯依据[9]。从存证方式看，常见方案主要包括完整数据上链和摘要数据上链两类。完整数据上链能够使数据直接成为链上记录，便于后续查询和验证，但在日志、文件、业务数据等规模较大的场景中，会带来较高的存储成本和处理开销，也可能导致敏感信息暴露。",

    "摘要数据上链则将原始数据保存在链下，只将哈希摘要、业务标识、提交时间和提交者等元数据写入链上。该方式更适合日志审计场景，因为日志数量多、内容可能包含路径和错误信息，直接上链会增加存储成本和隐私风险。需要注意的是，摘要上链只能证明某个摘要被登记，不能单独保证采集过程完整，因此还需要结合链下数据库、自动采集和审计比对机制。本文正是在这一思路下，将 SQLite 链下存储与 LogRegistry 链上哈希存证结合起来。":
    "摘要数据上链则将原始数据保存在链下，只将哈希摘要、业务标识、提交时间和提交者等元数据写入链上。该方式更适合日志审计场景，因为日志数量多、内容可能包含路径和错误信息，直接上链会增加存储成本和隐私风险。链上链下结合的日志安全存储研究说明，将日志原文与摘要证据分层保存是日志完整性保护中的可行思路[10]。需要注意的是，摘要上链只能证明某个摘要被登记，不能单独保证采集过程完整，因此还需要结合链下数据库、自动采集和审计比对机制。本文正是在这一思路下，将 SQLite 链下存储与 LogRegistry 链上哈希存证结合起来。",

    "日志审计是信息系统安全管理和运行维护中的重要内容。已有研究和工程实践通常围绕日志采集、集中存储、查询分析、异常检测和审计追溯展开。通过统一收集服务器、应用程序和任务执行过程中的日志，系统可以在发生异常时定位故障原因，在发生安全事件时还原操作过程，并为责任追溯提供依据。因此，日志审计不仅关注日志是否能够被记录，还关注日志是否完整、可信和可追溯。传统日志审计系统通常采用集中式日志平台或数据库保存日志数据。该方式便于统一检索、聚合分析和可视化展示，但在完整性保护方面仍存在不足。":
    "日志审计是信息系统安全管理和运行维护中的重要内容。已有研究和工程实践通常围绕日志采集、集中存储、查询分析、异常检测和审计追溯展开。防篡改审计日志研究强调日志完整性保护对事后取证和责任追溯的重要性[11]。通过统一收集服务器、应用程序和任务执行过程中的日志，系统可以在发生异常时定位故障原因，在发生安全事件时还原操作过程，并为责任追溯提供依据。因此，日志审计不仅关注日志是否能够被记录，还关注日志是否完整、可信和可追溯。安全日志研究通常关注日志生成、存储和验证阶段的完整性保护机制[12]。传统日志审计系统通常采用集中式日志平台或数据库保存日志数据。该方式便于统一检索、聚合分析和可视化展示，但在完整性保护方面仍存在不足。区块链防篡改日志研究表明，链上记录可增强日志事后验证与篡改识别能力[13]。",

    "由于日志原文和审计依据往往保存在同一管理域内，当系统权限被滥用或数据库被攻击时，日志可能被修改、删除或伪造。在完整性校验方面，哈希摘要、数字签名、时间戳和访问控制等技术常被用于增强日志可信性。此外，传统日志审计在自动化和闭环能力方面也存在不足。本课题面向上述问题，设计 Agent 自动采集、三方哈希比对和 hash_mismatch 告警机制，以提高日志审计过程的完整性和可观察性。":
    "由于日志原文和审计依据往往保存在同一管理域内，当系统权限被滥用或数据库被攻击时，日志可能被修改、删除或伪造。在完整性校验方面，哈希摘要、数字签名、时间戳和访问控制等技术常被用于增强日志可信性。此外，传统日志审计在自动化和闭环能力方面也存在不足；自动化日志分析研究覆盖日志解析、异常检测和可靠性分析等方向，但本文实现重点仍是哈希完整性审计[14]。深度学习日志异常检测属于相关研究方向，本文不实现该类语义检测，而聚焦日志内容是否与存证时一致[15]。本课题面向上述问题，设计 Agent 自动采集、三方哈希比对和 hash_mismatch 告警机制，以提高日志审计过程的完整性和可观察性。",

    "智能合约能够将预先定义的业务规则部署到区块链网络中，并按照合约逻辑自动执行。由于合约状态和交易记录可以在链上追溯，智能合约在可信审计场景中常用于固化写入规则、限定操作权限、保存证据摘要、提供查询接口和记录事件信息。相比仅依靠链下程序执行审计逻辑，智能合约可以为关键存证行为提供更加稳定的链上规则约束。在数据存证与审计场景中，智能合约通常承担证据登记和查询功能。例如，系统可以通过合约方法写入数据哈希、业务标识、提交时间和提交者地址，并通过事件记录写入结果。":
    "智能合约能够将预先定义的业务规则部署到区块链网络中，并按照合约逻辑自动执行。智能合约可将链上规则、状态和事件组织为可验证的程序化执行机制[16]。由于合约状态和交易记录可以在链上追溯，智能合约在可信审计场景中常用于固化写入规则、限定操作权限、保存证据摘要、提供查询接口和记录事件信息。智能合约能够把预先定义的规则部署到区块链上，并通过链上调用和状态记录提供可追溯执行依据[17]。相比仅依靠链下程序执行审计逻辑，智能合约可以为关键存证行为提供更加稳定的链上规则约束。在数据存证与审计场景中，智能合约通常承担证据登记和查询功能。例如，系统可以通过合约方法写入数据哈希、业务标识、提交时间和提交者地址，并通过事件记录写入结果。",

    "后续审计时，链下系统可以根据业务标识查询链上记录，再与本地数据库记录和重新计算得到的哈希进行比对。通过这种方式，智能合约成为链下业务数据与链上可信证据之间的连接层。本系统中的 LogRegistry 合约采用 AccessControl 管理 LOGGER_ROLE，只有被授权地址才能调用 storeLog 写入日志哈希。与此同时，智能合约也存在部署后修改成本较高、链上存储开销较大等限制，因此本文仅将日志哈希和必要元数据写入链上。":
    "后续审计时，链下系统可以根据业务标识查询链上记录，再与本地数据库记录和重新计算得到的哈希进行比对。通过这种方式，智能合约成为链下业务数据与链上可信证据之间的连接层。本系统中的 LogRegistry 合约采用 AccessControl 管理 LOGGER_ROLE，只有被授权地址才能调用 storeLog 写入日志哈希。与此同时，智能合约在固化规则的同时也存在攻击面和权限配置风险，因此权限边界和输入约束需要在设计阶段明确[18]。合约部署后修改成本较高，链上存储开销也较大，因此本文仅将日志哈希和必要元数据写入链上。",

    "区块链是一种以区块为基本组织单位、通过密码学机制和共识机制维护数据一致性的分布式账本技术。与传统中心化数据库相比，区块链不依赖单一节点保存全部可信依据，而是由多个节点共同维护账本状态。交易被确认并写入区块后，会与区块时间、交易内容、前一区块哈希等信息形成关联，后续区块继续在此基础上扩展。若试图修改历史数据，相关哈希值和区块链接关系也会发生变化，因此历史记录具有较强的不可篡改性和可追溯性。在日志审计场景中，区块链可以为链下日志提供独立的可信证明。":
    "区块链是一种以区块为基本组织单位、通过密码学机制和共识机制维护数据一致性的分布式账本技术[8]。与传统中心化数据库相比，区块链不依赖单一节点保存全部可信依据，而是由多个节点共同维护账本状态。交易被确认并写入区块后，会与区块时间、交易内容、前一区块哈希等信息形成关联，后续区块继续在此基础上扩展。若试图修改历史数据，相关哈希值和区块链接关系也会发生变化，因此历史记录具有较强的不可篡改性和可追溯性。在日志审计场景中，区块链可以为链下日志提供独立的可信证明。",

    "访问控制是智能合约安全设计中的关键环节。在日志存证场景中，若任意地址都能写入日志哈希，攻击者可能提交伪造摘要，进而削弱链上证据的可信性。OpenZeppelin AccessControl 提供基于角色的授权机制，可通过 DEFAULT_ADMIN_ROLE 管理角色，通过 LOGGER_ROLE 限定日志写入主体。与简单自定义权限字段相比，该机制接口清晰、可复用性较好，适合 LogRegistry 合约的写入权限控制。":
    "访问控制是智能合约安全设计中的关键环节。在日志存证场景中，若任意地址都能写入日志哈希，攻击者可能提交伪造摘要，进而削弱链上证据的可信性。基于角色的访问控制通过角色约束操作主体，适合解释日志哈希写入权限的边界[19]。OpenZeppelin `AccessControl` 提供角色授权机制，可用于实现 `DEFAULT_ADMIN_ROLE` 和 `LOGGER_ROLE` 等合约角色[20]。与简单自定义权限字段相比，该机制接口清晰、可复用性较好，适合 LogRegistry 合约的写入权限控制。",

    "哈希摘要技术是日志完整性校验的重要基础。哈希函数能够将任意长度的输入数据映射为固定长度的摘要值，该摘要值通常以十六进制字符串形式表示。对于同一份输入数据，使用相同哈希算法得到的结果应保持一致；安全哈希函数通常具有单向性、抗碰撞性和雪崩效应等特点。SHA-256 是安全哈希算法家族中的常用算法之一，其输出长度为 256 位，通常表示为固定长度的十六进制字符串。系统在日志首次提交时对日志正文计算哈希，并将该哈希作为提交时的 expectedHash；后续审计时，再对当前保存的日志正文重新计算 actualHash。":
    "哈希摘要技术是日志完整性校验的重要基础。哈希函数能够将任意长度的输入数据映射为固定长度的摘要值，该摘要值通常以十六进制字符串形式表示。对于同一份输入数据，使用相同哈希算法得到的结果应保持一致；安全哈希函数通常具有单向性、抗碰撞性和雪崩效应等特点。SHA-256 属于安全哈希标准中的常用算法，其输出长度为 256 位，可用于生成固定长度摘要[21]。系统在日志首次提交时对日志正文计算哈希，并将该哈希作为提交时的 expectedHash；后续审计时，再对当前保存的日志正文重新计算 actualHash。",

    "哈希摘要用于完整性校验时，重点并不在于还原日志原文，而在于比较同一输入在不同时刻得到的摘要是否一致。图 2-2 展示了日志内容变化导致哈希摘要变化的基本原理。":
    "哈希摘要用于完整性校验时，重点并不在于还原日志原文，而在于比较同一输入在不同时刻得到的摘要是否一致。哈希树思想体现了利用哈希结构组织数据完整性证明的经典方法[22]。图 2-2 展示了日志内容变化导致哈希摘要变化的基本原理。",

    "日志数据是系统运行过程中形成的证据材料，可用于安全审计、故障排查、责任追溯和运行状态分析。传统日志管理通常将日志原文保存在服务器文件系统、关系型数据库或集中日志平台中，虽然便于查询，但审计依据仍主要依赖中心化存储。本系统采用“链下保存日志原文、链上保存日志哈希摘要”的混合存储模型：日志进入系统后，后端服务先对日志原文计算 SHA-256 哈希值，再将哈希摘要及任务标识等必要元数据写入 LogRegistry 智能合约。":
    "日志数据是系统运行过程中形成的证据材料，可用于安全审计、故障排查、责任追溯和运行状态分析。传统日志管理通常将日志原文保存在服务器文件系统、关系型数据库或集中日志平台中，虽然便于查询，但审计依据仍主要依赖中心化存储。本系统采用“链下保存日志原文、链上保存日志哈希摘要”的混合存储模型：日志进入系统后，后端服务先对日志原文计算 SHA-256 哈希值，再将哈希摘要及任务标识等必要元数据写入 LogRegistry 智能合约。数据溯源研究表明，链上记录可作为链下数据状态和追溯路径的可信参照[9]。",

    "日志原文通常数量多、长度不固定且变化频繁，直接上链会增加存储开销，也可能暴露路径、错误信息等敏感内容。采用摘要上链后，SQLite 保存可查询的日志原文和业务记录，LogRegistry 保存可验证的哈希证据，两者分别承担业务存储和可信校验职责。图 3-1 展示了链下数据表与链上 LogRegistry 之间的对应关系。":
    "日志原文通常数量多、长度不固定且变化频繁，直接上链会增加存储开销，也可能暴露路径、错误信息等敏感内容。采用摘要上链后，SQLite 保存可查询的日志原文和业务记录，LogRegistry 保存可验证的哈希证据，两者分别承担业务存储和可信校验职责。链上链下结合的存储方式能够在保留链下数据查询能力的同时，将关键摘要写入链上作为完整性校验依据[10]。图 3-1 展示了链下数据表与链上 LogRegistry 之间的对应关系。",

    "本系统基于 OpenZeppelin AccessControl 实现角色控制。DEFAULT_ADMIN_ROLE 用于管理角色授权，LOGGER_ROLE 用于控制日志哈希写入权限。后端服务使用具备 LOGGER_ROLE 的账户调用 storeLog，将日志哈希写入链上；未获得该角色的账户即使知道合约地址和方法，也不能完成有效写入。":
    "本系统基于 OpenZeppelin AccessControl 实现角色控制。DEFAULT_ADMIN_ROLE 用于管理角色授权，LOGGER_ROLE 用于控制日志哈希写入权限。本系统使用 `AccessControl` 将链上写入行为限定为具备 `LOGGER_ROLE` 的地址，以降低伪造写入对链上证据可信性的影响[20]。后端服务使用具备 LOGGER_ROLE 的账户调用 storeLog，将日志哈希写入链上；未获得该角色的账户即使知道合约地址和方法，也不能完成有效写入。",

    "与自定义权限字段相比，AccessControl 提供成熟的角色管理接口和事件记录，便于授权、撤销和审计。该设计将链下业务系统的提交行为与链上权限边界对应起来，使 LogRegistry 合约不仅保存证据摘要，也能够限制可信写入来源。":
    "与自定义权限字段相比，AccessControl 提供成熟的角色管理接口和事件记录，便于授权、撤销和审计。该设计也回应了智能合约权限配置错误可能削弱系统安全边界的问题[18]，将链下业务系统的提交行为与链上权限边界对应起来，使 LogRegistry 合约不仅保存证据摘要，也能够限制可信写入来源。",

    "三方哈希比对是本系统审计机制的核心。系统在日志提交阶段计算日志内容的 SHA-256 摘要，并将该值写入 log_hash_records，作为 expectedHash；同时将对应哈希写入 LogRegistry 合约，形成链上 onChainHash。审计阶段，系统再次读取 logs 表中的日志原文并重新计算 actualHash。":
    "三方哈希比对是本系统审计机制的核心。系统在日志提交阶段计算日志内容的 SHA-256 摘要，并将该值写入 log_hash_records，作为 expectedHash[21]；同时将对应哈希写入 LogRegistry 合约，形成链上 onChainHash。审计阶段，系统再次读取 logs 表中的日志原文并重新计算 actualHash。",

    "审计逻辑以 expectedHash、actualHash 和 onChainHash 的一致性为判断依据。若三者完全一致，说明当前日志原文、数据库存证记录和链上存证记录相互匹配，审计结果为 passed。若 actualHash 与 expectedHash 不一致，说明链下日志原文可能被修改；若 expectedHash 与 onChainHash 不一致，则说明数据库记录与链上存证存在差异，审计结果为 failed。":
    "审计逻辑以 expectedHash、actualHash 和 onChainHash 的一致性为判断依据。若三者完全一致，说明当前日志原文、数据库存证记录和链上存证记录相互匹配，审计结果为 passed。若 actualHash 与 expectedHash 不一致，说明链下日志原文可能被修改；该判断体现了安全日志研究中通过完整性校验识别事后篡改的基本思想[11]。若 expectedHash 与 onChainHash 不一致，则说明数据库记录与链上存证存在差异，审计结果为 failed。",
}


REFERENCES = [
    "[1] KENT K, SOUPPAYA M P. Guide to computer security log management[R]. Gaithersburg: National Institute of Standards and Technology, 2006. DOI: 10.6028/NIST.SP.800-92.",
    "[2] 韩菊茹, 纪兆轩, 李一鸣, 等. 基于区块链的可信日志存储与验证系统[J]. 计算机工程, 2019, 45(5): 13-17. DOI: 10.19678/j.issn.1000-3428.0053385.",
    "[3] 袁勇, 王飞跃. 区块链技术发展现状与展望[J]. 自动化学报, 2016, 42(4): 481-494. DOI: 10.16383/j.aas.2016.c160158.",
    "[4] 钱卫宁, 邵奇峰, 朱燕超, 等. 区块链与可信数据管理:问题与方法[J]. 软件学报, 2018, 29(1): 150-159. DOI: 10.13328/j.cnki.jos.005434.",
    "[5] 刘敖迪, 杜学绘, 王娜, 等. 区块链技术及其在信息安全领域的研究进展[J]. 软件学报, 2018, 29(7): 2092-2115. DOI: 10.13328/j.cnki.jos.005589.",
    "[6] 范伟, 李海波, 张珠君. 区块链数字取证:技术及架构研究[J]. 通信学报, 2024, 45(12): 124-141. DOI: 10.11959/j.issn.1000-436x.2024204.",
    "[7] LEMIEUX V L. Trusting records: is Blockchain technology the answer?[J]. Records Management Journal, 2016, 26(2): 110-139. DOI: 10.1108/RMJ-12-2015-0042.",
    "[8] YAGA D, MELL P, ROBY N, et al. Blockchain technology overview[R]. Gaithersburg: National Institute of Standards and Technology, 2018. DOI: 10.6028/NIST.IR.8202.",
    "[9] LIANG X, SHETTY S, TOSH D, et al. ProvChain: a blockchain-based data provenance architecture in cloud environment with enhanced privacy and availability[C]//2017 17th IEEE/ACM International Symposium on Cluster, Cloud and Grid Computing (CCGRID). Piscataway: IEEE, 2017: 468-477. DOI: 10.1109/CCGRID.2017.8.",
    "[10] 吕建富, 赖英旭, 刘静. 基于链上链下相结合的日志安全存储与检索[J]. 计算机科学, 2020, 47(3): 298-303. DOI: 10.11896/jsjkx.190200298.",
    "[11] SCHNEIER B, KELSEY J. Secure audit logs to support computer forensics[J]. ACM Transactions on Information and System Security, 1999, 2(2): 159-176. DOI: 10.1145/317087.317089.",
    "[12] MA D, TSUDIK G. A new approach to secure logging[J]. ACM Transactions on Storage, 2009, 5(1): 1-21. DOI: 10.1145/1502777.1502779.",
    "[13] SHEKHTMAN L, WAISBARD E. EngraveChain: a blockchain-based tamper-proof distributed log system[J]. Future Internet, 2021, 13(6): 143. DOI: 10.3390/fi13060143.",
    "[14] HE S, HE P, CHEN Z, et al. A survey on automated log analysis for reliability engineering[J]. ACM Computing Surveys, 2022, 54(6): 1-37. DOI: 10.1145/3460345.",
    "[15] DU M, LI F, ZHENG G, et al. DeepLog: anomaly detection and diagnosis from system logs through deep learning[C]//Proceedings of the 2017 ACM SIGSAC Conference on Computer and Communications Security. New York: ACM, 2017: 1285-1298. DOI: 10.1145/3133956.3134015.",
    "[16] 贺海武, 延安, 陈泽华. 基于区块链的智能合约技术与应用综述[J]. 计算机研究与发展, 2018, 55(11): 2452-2466.",
    "[17] CHRISTIDIS K, DEVETSIKIOTIS M. Blockchains and smart contracts for the Internet of Things[J]. IEEE Access, 2016, 4: 2292-2303. DOI: 10.1109/ACCESS.2016.2566339.",
    "[18] ATZEI N, BARTOLETTI M, CIMOLI T. A survey of attacks on Ethereum smart contracts (SoK)[C]//Principles of Security and Trust. Berlin: Springer, 2017: 164-186. DOI: 10.1007/978-3-662-54455-6_8.",
    "[19] SANDHU R S, COYNE E J, FEINSTEIN H L, et al. Role-based access control models[J]. Computer, 1996, 29(2): 38-47. DOI: 10.1109/2.485845.",
    "[20] OpenZeppelin. Access Control: OpenZeppelin Contracts 5.x[EB/OL]. [2026-05-06]. https://docs.openzeppelin.com/contracts/5.x/access-control.",
    "[21] National Institute of Standards and Technology. FIPS PUB 180-4 Secure hash standard[S]. Gaithersburg: National Institute of Standards and Technology, 2015. DOI: 10.6028/NIST.FIPS.180-4.",
    "[22] MERKLE R C. A digital signature based on a conventional encryption function[C]//Advances in Cryptology — CRYPTO '87. Berlin: Springer, 1988: 369-378. DOI: 10.1007/3-540-48184-2_32.",
]


def all_text(doc: Document) -> str:
    return "\n".join(p.text for p in doc.paragraphs)


def count_captions(doc: Document, prefix: str) -> int:
    return sum(1 for p in doc.paragraphs if p.text.strip().startswith(prefix))


def doc_metrics(doc: Document) -> dict[str, int]:
    return {
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "inline_shapes": len(doc.inline_shapes),
        "figure_captions": count_captions(doc, "图 "),
        "table_captions": count_captions(doc, "表 "),
    }


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def set_paragraph_text_with_superscript_citations(paragraph, text: str) -> None:
    paragraph.clear()
    for part in re.split(r"(\[\d+\])", text):
        if not part:
            continue
        run = paragraph.add_run(part)
        if re.fullmatch(r"\[\d+\]", part):
            run.font.superscript = True


def remove_existing_reference_section(doc: Document) -> bool:
    removed = False
    while True:
        paragraphs = list(doc.paragraphs)
        start = next((i for i, p in enumerate(paragraphs) if p.text.strip() == "参考文献"), None)
        if start is None:
            return removed
        end = len(paragraphs)
        for i in range(start + 1, len(paragraphs)):
            if paragraphs[i].text.strip() in {"致谢", "附录", "外文原文及译文"}:
                end = i
                break
        for p in paragraphs[start:end]:
            remove_paragraph(p)
        removed = True


def insert_references_before_acknowledgement(doc: Document) -> str:
    paragraphs = list(doc.paragraphs)
    ack_para = next((p for p in paragraphs if p.text.strip() == "致谢"), None)
    if ack_para is None:
        doc.add_paragraph()
        doc.add_paragraph("参考文献")
        for ref in REFERENCES:
            doc.add_paragraph(ref)
        return "未找到“致谢”，已在文档末尾追加“参考文献”。"

    ack_para.insert_paragraph_before("")
    heading = ack_para.insert_paragraph_before("参考文献")
    heading.style = ack_para.style
    for ref in REFERENCES:
        ack_para.insert_paragraph_before(ref)
    ack_para.insert_paragraph_before("")
    return "已在第 6 章之后、致谢之前新增“参考文献”。"


def find_reference_range(doc: Document) -> tuple[int | None, int | None]:
    paragraphs = list(doc.paragraphs)
    start = next((i for i, p in enumerate(paragraphs) if p.text.strip() == "参考文献"), None)
    if start is None:
        return None, None
    end = len(paragraphs)
    for i in range(start + 1, len(paragraphs)):
        if paragraphs[i].text.strip() in {"致谢", "附录", "外文原文及译文"}:
            end = i
            break
    return start, end


def collect_body_citations(doc: Document) -> list[int]:
    start, _ = find_reference_range(doc)
    limit = start if start is not None else len(doc.paragraphs)
    citations: list[int] = []
    for p in list(doc.paragraphs)[:limit]:
        citations.extend(int(m.group(1)) for m in re.finditer(r"\[(\d+)\]", p.text))
    return citations


def collect_reference_numbers(doc: Document) -> list[int]:
    start, end = find_reference_range(doc)
    if start is None or end is None:
        return []
    nums: list[int] = []
    for p in list(doc.paragraphs)[start + 1:end]:
        m = re.match(r"^\[(\d+)\]\s+", p.text.strip())
        if m:
            nums.append(int(m.group(1)))
    return nums


def collect_superscript_status(doc: Document) -> tuple[int, int, list[str]]:
    start, _ = find_reference_range(doc)
    limit = start if start is not None else len(doc.paragraphs)
    total = 0
    superscript = 0
    non_superscript_samples: list[str] = []
    for p in list(doc.paragraphs)[:limit]:
        for run in p.runs:
            for match in re.finditer(r"\[\d+\]", run.text):
                total += 1
                if run.font.superscript is True:
                    superscript += 1
                elif len(non_superscript_samples) < 5:
                    non_superscript_samples.append(match.group(0))
    return total, superscript, non_superscript_samples


def is_continuous(nums: list[int]) -> bool:
    if not nums:
        return False
    return sorted(set(nums)) == list(range(1, max(nums) + 1))


def language_counts() -> tuple[int, int]:
    chinese_numbers = {2, 3, 4, 5, 6, 10, 16}
    return len(chinese_numbers), len(REFERENCES) - len(chinese_numbers)


def make_report(
    *,
    backup_path: Path,
    clean_source_path: Path,
    before_current_metrics: dict[str, int],
    clean_source_metrics: dict[str, int],
    after_metrics: dict[str, int],
    replaced_count: int,
    missing_replacements: list[str],
    reference_location: str,
    removed_existing_refs: bool,
    doc: Document,
) -> str:
    text = all_text(doc)
    citations = collect_body_citations(doc)
    reference_nums = collect_reference_numbers(doc)
    citation_unique = sorted(set(citations))
    reference_unique = sorted(set(reference_nums))
    citation_counter = Counter(citations)
    reference_counter = Counter(reference_nums)
    missing_in_refs = [n for n in citation_unique if n not in reference_unique]
    uncited_refs = [n for n in reference_unique if n not in citation_unique]
    duplicate_ref_nums = sorted(n for n, c in reference_counter.items() if c > 1)
    citation_total, citation_superscript, non_sup_samples = collect_superscript_status(doc)
    chinese_count, foreign_count = language_counts()

    expected_data = [
        "107.03 ms",
        "9.33 条/秒",
        "3067.77 ms",
        "15659.44 ms",
        "35032.13 ms",
        "auditStatus 为 failed",
        "alertGenerated 为 true",
    ]
    key_terms = [
        "LogRegistry",
        "LOGGER_ROLE",
        "provider.getCode(address)",
        "contract_address",
        "expectedHash",
        "actualHash",
        "onChainHash",
        "hash_mismatch",
    ]
    data_status = {item: (item in text) for item in expected_data}
    term_counts = {term: text.count(term) for term in key_terms}

    body_continuous = sorted(set(citations)) == list(range(1, len(REFERENCES) + 1))
    ref_continuous = is_continuous(reference_nums) and max(reference_unique or [0]) == len(REFERENCES)
    all_citations_have_refs = not missing_in_refs
    no_uncited_refs = not uncited_refs
    all_superscript = citation_total > 0 and citation_total == citation_superscript
    media_preserved = (
        clean_source_metrics["tables"] == after_metrics["tables"]
        and clean_source_metrics["inline_shapes"] == after_metrics["inline_shapes"]
        and clean_source_metrics["figure_captions"] == after_metrics["figure_captions"]
        and clean_source_metrics["table_captions"] == after_metrics["table_captions"]
    )

    lines = [
        "# 论文参考文献插入校验报告",
        "",
        f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        "",
        "## 一、处理结果",
        "",
        f"- Word 文件：`{DOCX_PATH}`",
        f"- 修改前备份文件：`{backup_path}`",
        f"- 使用的干净基础文档：`{clean_source_path}`",
        f"- 参考文献章节位置：{reference_location}",
        f"- 是否移除基础文档旧参考文献章节：{'是' if removed_existing_refs else '否，基础文档未发现旧参考文献章节'}",
        f"- 正文段落替换数量：{replaced_count}",
        f"- 未匹配替换段落数量：{len(missing_replacements)}",
        f"- 最终参考文献构成：中文 {chinese_count} 篇，外文 {foreign_count} 篇",
        "",
        "## 二、引用与参考文献校验",
        "",
        f"- 正文引用总数：{len(citations)}",
        f"- 正文引用唯一编号数：{len(citation_unique)}",
        f"- 参考文献总数：{len(reference_nums)}",
        f"- 正文引用编号是否连续：{'是' if body_continuous else '否'}",
        f"- 参考文献编号是否连续：{'是' if ref_continuous else '否'}",
        f"- 每个引用编号是否在参考文献列表中存在：{'是' if all_citations_have_refs else '否'}",
        f"- 是否存在未被正文引用的参考文献：{'否' if no_uncited_refs else '是，编号为 ' + ', '.join(map(str, uncited_refs))}",
        f"- 是否存在重复参考文献编号：{'否' if not duplicate_ref_nums else '是，编号为 ' + ', '.join(map(str, duplicate_ref_nums))}",
        f"- 正文引用是否设置为上标：{'是' if all_superscript else '否'}",
        f"- 正文上标引用数量：{citation_superscript}/{citation_total}",
        "",
        "## 三、编号对应检查",
        "",
        "| 编号 | 正文出现次数 | 参考文献是否存在 |",
        "| --- | ---: | --- |",
    ]
    for n in range(1, len(REFERENCES) + 1):
        lines.append(f"| [{n}] | {citation_counter.get(n, 0)} | {'是' if n in reference_unique else '否'} |")

    lines.extend([
        "",
        "## 四、实验数据与关键术语校验",
        "",
        "| 检查项 | 是否保留 |",
        "| --- | --- |",
    ])
    for item, ok in data_status.items():
        lines.append(f"| `{item}` | {'是' if ok else '否'} |")
    lines.extend([
        "",
        "| 关键术语 | 当前出现次数 | 是否保留 |",
        "| --- | ---: | --- |",
    ])
    for term, count in term_counts.items():
        lines.append(f"| `{term}` | {count} | {'是' if count > 0 else '否'} |")

    lines.extend([
        "",
        "## 五、图片、表格和图题校验",
        "",
        "| 项目 | 当前 Word 修改前 | 干净基础文档 | 修改后 | 是否相对基础文档一致 |",
        "| --- | ---: | ---: | ---: | --- |",
        f"| 表格数量 | {before_current_metrics['tables']} | {clean_source_metrics['tables']} | {after_metrics['tables']} | {'是' if clean_source_metrics['tables'] == after_metrics['tables'] else '否'} |",
        f"| 内嵌图片数量 | {before_current_metrics['inline_shapes']} | {clean_source_metrics['inline_shapes']} | {after_metrics['inline_shapes']} | {'是' if clean_source_metrics['inline_shapes'] == after_metrics['inline_shapes'] else '否'} |",
        f"| 图题数量 | {before_current_metrics['figure_captions']} | {clean_source_metrics['figure_captions']} | {after_metrics['figure_captions']} | {'是' if clean_source_metrics['figure_captions'] == after_metrics['figure_captions'] else '否'} |",
        f"| 表题数量 | {before_current_metrics['table_captions']} | {clean_source_metrics['table_captions']} | {after_metrics['table_captions']} | {'是' if clean_source_metrics['table_captions'] == after_metrics['table_captions'] else '否'} |",
        "",
        f"是否保留已有图片、表格和图题：{'是' if media_preserved else '否'}",
        "",
        "## 六、风险提示",
        "",
    ])
    if missing_replacements:
        lines.append("以下原段落未匹配到，可能说明基础文档正文已被修改：")
        for item in missing_replacements:
            lines.append(f"- {item[:80]}...")
    else:
        lines.append("- 所有计划替换段落均已匹配并处理。")
    if non_sup_samples:
        lines.append(f"- 存在未上标的正文引用样例：{', '.join(non_sup_samples)}。")
    else:
        lines.append("- 正文引用 run 均已设置为 Word 上标。")
    lines.append("- 第 4、5、6 章未新增外部引用，实验数据仍以项目真实结果为准。")
    lines.append("- 当前 Word 中上一轮英文参考文献已通过使用干净基础文档重做的方式移除。")

    return "\n".join(lines) + "\n"


def main() -> None:
    if not DOCX_PATH.exists():
        raise FileNotFoundError(DOCX_PATH)
    if not CLEAN_SOURCE_PATH.exists():
        raise FileNotFoundError(CLEAN_SOURCE_PATH)

    timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    backup_path = DOCX_PATH.parent / f"Graduation-thesis-before-reference-insert-{timestamp}.docx"
    shutil.copy2(DOCX_PATH, backup_path)

    before_current_doc = Document(str(DOCX_PATH))
    clean_doc = Document(str(CLEAN_SOURCE_PATH))
    before_current_metrics = doc_metrics(before_current_doc)
    clean_source_metrics = doc_metrics(clean_doc)

    doc = clean_doc
    removed_existing_refs = remove_existing_reference_section(doc)

    replaced_count = 0
    matched_keys: set[str] = set()
    for paragraph in doc.paragraphs:
        original = paragraph.text
        if original in REPLACEMENTS:
            set_paragraph_text_with_superscript_citations(paragraph, REPLACEMENTS[original])
            replaced_count += 1
            matched_keys.add(original)

    missing_replacements = [key for key in REPLACEMENTS if key not in matched_keys]
    reference_location = insert_references_before_acknowledgement(doc)
    doc.save(str(DOCX_PATH))

    after_doc = Document(str(DOCX_PATH))
    after_metrics = doc_metrics(after_doc)
    report = make_report(
        backup_path=backup_path,
        clean_source_path=CLEAN_SOURCE_PATH,
        before_current_metrics=before_current_metrics,
        clean_source_metrics=clean_source_metrics,
        after_metrics=after_metrics,
        replaced_count=replaced_count,
        missing_replacements=missing_replacements,
        reference_location=reference_location,
        removed_existing_refs=removed_existing_refs,
        doc=after_doc,
    )
    REPORT_PATH.write_text(report, encoding="utf-8")

    print(f"backup: {backup_path}")
    print(f"saved: {DOCX_PATH}")
    print(f"report: {REPORT_PATH}")
    print(f"replaced: {replaced_count}/{len(REPLACEMENTS)}")


if __name__ == "__main__":
    main()
