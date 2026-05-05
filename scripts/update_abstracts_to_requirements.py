from __future__ import annotations

import re
from pathlib import Path

from docx import Document


DOCX_PATH = Path("output/Graduation-thesis.docx")
FALLBACK_PATH = Path("output/Graduation-thesis-abstracts-updated.docx")


CHINESE_ABSTRACT = (
    "随着信息系统在业务处理、自动化运维和任务执行场景中的广泛应用，日志数据已成为安全审计、故障排查和责任追溯的重要依据。传统中心化日志通常将日志原文、哈希记录和审计结果保存在同一管理域内，存在被篡改、删除或伪造后难以及时验证的问题。针对上述问题，本文设计并实现了一种基于区块链的可信任务日志审计系统。系统采用“链下保存日志原文、链上保存日志哈希摘要”的混合存储模型，由 Agent 完成日志增量采集，Server 负责日志入库、SHA-256 哈希计算和 LogRegistry 智能合约调用，SQLite 保存日志、存证、审计和告警记录，Web 前端展示日志、审计和告警结果。审计阶段，系统重新计算日志原文得到 actualHash，并与数据库中的 expectedHash、链上的 onChainHash 进行三方比对；当哈希不一致时，生成 failed 审计记录和 hash_mismatch 告警。系统还针对本地链重启和多次部署场景，引入合约代码存在性校验与历史合约地址回溯策略，以减少空地址和历史记录误判。实验结果表明，LogRegistry 合约测试 8 项全部通过；日志批量提交 100 次成功率为 100%，平均响应时间约 107.03 ms，吞吐量约 9.33 条/秒；批量审计 100、500、1000 条数据 5 轮均执行成功；篡改检测实验能够识别日志内容被修改并生成告警。"
)


ENGLISH_ABSTRACT = (
    "With the widespread use of information systems in business processing, automated operation and maintenance, and task execution scenarios, log data has become an important basis for security auditing, fault diagnosis, and responsibility tracing. Traditional centralized log management usually stores original logs, hash records, and audit results within the same management domain. Once logs are tampered with, deleted, or forged, it is difficult to provide independent and reliable evidence for later verification. To address this problem, this thesis designs and implements a blockchain-based trusted task log auditing system. The system adopts a hybrid storage model in which original logs are stored off-chain and log hash digests are stored on-chain. The Agent performs incremental log collection, the Server stores logs in SQLite, calculates SHA-256 hashes, and calls the LogRegistry smart contract, while SQLite maintains log records, evidence records, audit records, and alert records. The Web frontend displays logs, audit results, and alert information. During auditing, actualHash is recalculated from the current log content and compared with expectedHash stored in the database and onChainHash queried from the smart contract. If the three hashes are consistent, the audit result is passed; otherwise, a failed audit record and a hash_mismatch alert are generated. The system also introduces contract code existence verification and historical contract address tracing for local chain restart and repeated deployment scenarios. Experimental results show that all 8 LogRegistry contract tests passed, 100 batch log submissions achieved a 100% success rate, the average response time was about 107.03 ms, and the throughput was about 9.33 records per second. Batch audits for 100, 500, and 1000 records were successfully completed in five rounds, and the tamper detection experiment showed that modified log content could be identified and alerted."
)


def set_after_heading(doc: Document, heading: str, value: str) -> None:
    for index, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == heading:
            doc.paragraphs[index + 1].text = value
            return
    raise ValueError(f"未找到标题：{heading}")


def english_word_count(text: str) -> int:
    return len(re.findall(r"[A-Za-z]+(?:[-'][A-Za-z]+)?|\d+(?:\.\d+)?", text))


def main() -> None:
    doc = Document(DOCX_PATH)
    set_after_heading(doc, "摘要", CHINESE_ABSTRACT)
    set_after_heading(doc, "Abstract", ENGLISH_ABSTRACT)

    try:
        doc.save(DOCX_PATH)
        saved_path = DOCX_PATH
    except PermissionError:
        doc.save(FALLBACK_PATH)
        saved_path = FALLBACK_PATH

    chinese_cjk_count = len(re.findall(r"[\u4e00-\u9fff]", CHINESE_ABSTRACT))
    chinese_total_chars = len(CHINESE_ABSTRACT)
    english_words = english_word_count(ENGLISH_ABSTRACT)

    print(f"saved={saved_path}")
    print(f"chinese_cjk_count={chinese_cjk_count}")
    print(f"chinese_total_chars={chinese_total_chars}")
    print(f"english_words={english_words}")


if __name__ == "__main__":
    main()
