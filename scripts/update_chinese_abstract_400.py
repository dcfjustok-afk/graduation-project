from __future__ import annotations

from pathlib import Path

from docx import Document


DOCX_PATH = Path("output/Graduation-thesis.docx")
FALLBACK_PATH = Path("output/Graduation-thesis-abstract-400.docx")


ABSTRACT = (
    "随着信息系统在业务处理、自动化运维和任务执行场景中的广泛应用，日志数据已成为安全审计、故障排查和责任追溯的重要依据。针对传统中心化日志容易被篡改、删除或伪造，且审计结果缺少独立可信依据的问题，本文设计并实现了一种基于区块链的可信任务日志审计系统。系统采用“链下保存日志原文、链上保存日志哈希摘要”的混合存储模型，由 Agent 完成日志增量采集，Server 负责日志入库、SHA-256 哈希计算和 LogRegistry 智能合约调用，SQLite 保存日志、存证、审计和告警记录，Web 前端展示日志、审计和告警结果。审计阶段，系统重新计算日志原文得到 actualHash，并与数据库中的 expectedHash、链上的 onChainHash 进行三方比对；当哈希不一致时，生成 failed 审计记录和 hash_mismatch 告警。实验结果表明，LogRegistry 合约测试 8 项全部通过；日志批量提交 100 次成功率为 100%，平均响应时间约 107.03 ms，吞吐量约 9.33 条/秒；批量审计 100、500、1000 条数据均能完成；篡改检测实验中系统能够识别日志内容被修改并生成告警。"
)


def main() -> None:
    doc = Document(DOCX_PATH)
    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == "摘要":
            doc.paragraphs[idx + 1].text = ABSTRACT
            break
    else:
        raise ValueError("未找到摘要标题")
    try:
        doc.save(DOCX_PATH)
        print(f"saved={DOCX_PATH}")
    except PermissionError:
        doc.save(FALLBACK_PATH)
        print(f"target_locked_saved_fallback={FALLBACK_PATH}")
    print(f"abstract_chars={len(ABSTRACT)}")


if __name__ == "__main__":
    main()
