from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document


DOCX_PATH = Path("output/Graduation-thesis.docx")
SAFE_BASE = Path("output/Graduation-thesis-before-45page-prefix-rewrite.docx")
BACKUP_PATH = Path("output/Graduation-thesis-before-45page-rewrite-current.docx")


SECTION_TARGETS = {
    "摘要": 380,
    "Abstract": 1120,
    "1.1 研究背景及意义": 600,
    "1.2 国内外研究现状": 330,
    "1.2.1 区块链存证技术研究现状": 420,
    "1.2.2 日志审计与完整性校验研究现状": 430,
    "1.2.3 智能合约在可信审计中的应用现状": 430,
    "1.3 本文主要研究内容": 680,
    "1.4 论文组织结构": 170,
    "1.5 本章小结": 130,
    "2.1 区块链与链上存证技术": 430,
    "2.2 智能合约与访问控制技术": 520,
    "2.3 哈希摘要与日志完整性校验技术": 450,
    "2.4 系统开发相关技术": 100,
    "2.4.1 React 前端框架": 250,
    "2.4.2 Express 后端服务框架": 250,
    "2.4.3 SQLite 数据存储技术": 250,
    "2.4.4 Hardhat 智能合约开发环境": 170,
    "2.5 本章小结": 250,
    "3.1.1 链下日志存储与链上哈希存证模型": 420,
    "3.1.2 日志采集、存证、审计、告警闭环流程": 540,
    "3.2.1 日志标准化封装与任务标识设计": 420,
    "3.2.2 SHA-256 日志摘要生成方法": 410,
    "3.2.3 链上存证字段与链下记录映射关系": 620,
    "3.3.1 合约数据结构设计": 590,
    "3.3.2 基于 AccessControl 的写入权限控制": 560,
    "3.3.3 日志哈希写入与任务维度查询方法": 520,
    "3.4.1 数据库哈希、链上哈希与重算哈希三方比对": 700,
    "3.4.2 历史合约地址回溯审计策略": 480,
    "3.4.3 合约代码存在性校验机制": 520,
    "3.5.1 审计状态判定规则": 560,
    "3.5.2 哈希不一致告警生成流程": 470,
    "3.6 本章小结": 300,
    "4.2 系统需求分析": 220,
    "4.2.1 日志自动采集需求": 260,
    "4.2.2 链上可信存证需求": 270,
    "4.2.3 审计校验需求": 260,
    "4.2.4 异常告警需求": 270,
    "4.2.5 可视化展示需求": 270,
    "4.3 非功能需求分析": 100,
    "4.3.1 数据可信性需求": 270,
    "4.3.2 系统可用性需求": 270,
    "4.3.3 可维护性与可扩展性需求": 290,
    "4.4 系统总体架构设计": 700,
    "4.5 系统功能模块设计": 100,
    "4.5.1 日志采集 Agent 模块": 140,
    "4.5.2 后端服务模块": 220,
    "4.5.3 区块链存证模块": 140,
    "4.5.4 审计告警模块": 210,
    "4.5.5 前端可视化模块": 110,
    "4.6 系统数据流与业务流程设计": 680,
    "4.7 数据库设计": 150,
    "4.7.1 日志数据表设计": 280,
    "4.7.2 存证记录表设计": 300,
    "4.7.3 审计记录与告警表设计": 230,
    "4.7.4 Agent 状态表设计": 300,
    "4.8 本章小结": 250,
    "5.1 系统开发环境与工程结构": 620,
    "5.2 日志采集 Agent 实现": 200,
    "5.2.1 增量读取与偏移量持久化实现": 580,
    "5.2.2 失败重试与状态同步实现": 520,
    "5.3 后端服务与数据持久化实现": 190,
    "5.3.1 后端分层架构与核心接口实现": 460,
    "5.3.2 日志入库、哈希计算与上链流程实现": 500,
    "5.3.3 审计记录与告警数据表实现": 480,
    "5.4 前端可视化模块实现": 260,
    "5.4.1 系统总览与日志中心实现": 280,
    "5.4.2 审计管理与告警管理实现": 360,
    "5.5 功能测试与闭环验证": 150,
    "5.5.1 智能合约功能测试": 280,
    "5.5.2 后端接口与审计闭环测试": 330,
    "5.5.3 Agent 采集与前端展示测试": 390,
    "5.6 性能实验与篡改检测实验": 120,
    "5.6.1 日志批量提交实验": 360,
    "5.6.2 批量审计性能实验": 340,
    "5.6.3 篡改检测实验与结果分析": 430,
    "5.7 本章小结": 240,
    "6.1 研究工作总结": 500,
    "6.2 系统创新点": 580,
    "6.3 存在不足": 500,
    "6.4 后续优化方向": 500,
    "6.5 本章小结": 250,
    "致谢": 240,
}


PRIORITY_TERMS_BY_SECTION = {
    "3.2.3 链上存证字段与链下记录映射关系": [
        "logs",
        "log_hash_records",
        "LogRegistry",
        "log_id",
        "task_id",
        "log_hash",
        "contract_address",
        "transaction_hash",
        "block_number",
        "on_chain_status",
    ],
    "3.3.1 合约数据结构设计": [
        "LogRecord",
        "taskId",
        "logHash",
        "createdAt",
        "submitter",
        "records",
        "taskIdToRecordIds",
    ],
    "3.3.2 基于 AccessControl 的写入权限控制": [
        "LOGGER_ROLE",
        "DEFAULT_ADMIN_ROLE",
        "AccessControl",
        "OpenZeppelin",
    ],
    "3.4.1 数据库哈希、链上哈希与重算哈希三方比对": [
        "expectedHash",
        "actualHash",
        "onChainHash",
        "passed",
        "failed",
        "pending",
    ],
    "3.4.2 历史合约地址回溯审计策略": ["contract_address", "历史合约地址", "多次部署", "回溯"],
    "3.4.3 合约代码存在性校验机制": ["provider.getCode", "合约代码", "空地址", "字节码"],
    "3.5.1 审计状态判定规则": ["passed", "failed", "pending", "expectedHash", "actualHash", "onChainHash"],
    "3.5.2 哈希不一致告警生成流程": ["hash_mismatch", "alerts", "related_log_id", "related_audit_id"],
    "5.6.1 日志批量提交实验": ["100", "107.03", "9.33", "100%"],
    "5.6.2 批量审计性能实验": ["100", "500", "1000", "3067.77", "15659.44", "35032.13"],
    "5.6.3 篡改检测实验与结果分析": ["auditStatus", "failed", "alertGenerated", "true"],
    "6.2 系统创新点": ["三方哈希比对", "历史合约地址", "provider.getCode", "闭环"],
}


GLOBAL_PRIORITY_TERMS = [
    "LogRegistry",
    "LOGGER_ROLE",
    "provider.getCode",
    "contract_address",
    "expectedHash",
    "actualHash",
    "onChainHash",
    "hash_mismatch",
    "107.03",
    "9.33",
    "3067.77",
    "15659.44",
    "35032.13",
    "auditStatus",
    "alertGenerated",
]


def is_heading(text: str) -> bool:
    text = text.strip()
    if not text:
        return False
    if text in {"摘要", "Abstract", "致谢", "参考文献"}:
        return True
    if text.startswith("关键词：") or text.startswith("Keywords:"):
        return True
    if text.startswith("第 ") and "章" in text:
        return True
    parts = text.split(maxsplit=1)
    if not parts:
        return False
    marker = parts[0]
    return marker.count(".") >= 1 and all(part.isdigit() for part in marker.split("."))


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)
    paragraph._p = paragraph._element = None


def split_sentences(text: str) -> list[str]:
    text = re.sub(r"\s+", " ", text.strip())
    if not text:
        return []
    pieces = re.findall(r".+?(?:[。！？；;.!?]|$)", text)
    return [p.strip() for p in pieces if p.strip()]


def normalize_paragraph(text: str) -> str:
    text = re.sub(r"\s+", " ", text.strip())
    replacements = [
        ("在本系统中，", ""),
        ("从整体来看，", ""),
        ("需要指出的是，", ""),
        ("与此同时，", "同时，"),
        ("因此，", "因此，"),
        ("可以看出，", ""),
        ("进一步", ""),
        ("一定程度上", ""),
        ("较为", ""),
        ("相对", ""),
    ]
    for old, new in replacements:
        text = text.replace(old, new)
    return text.strip()


def group_sentences(sentences: list[str], max_len: int = 260) -> list[str]:
    paragraphs: list[str] = []
    current = ""
    for sentence in sentences:
        if not current:
            current = sentence
            continue
        if len(current) + len(sentence) <= max_len:
            current += sentence
        else:
            paragraphs.append(current)
            current = sentence
    if current:
        paragraphs.append(current)
    return paragraphs


def condense_text(section: str, paragraphs: list[str], target: int) -> list[str]:
    if sum(len(p) for p in paragraphs) <= target:
        return [normalize_paragraph(p) for p in paragraphs if p.strip()]

    captions: list[str] = []
    sentence_groups: list[list[str]] = []
    for paragraph in paragraphs:
        text = normalize_paragraph(paragraph)
        if not text:
            continue
        if text.startswith(("图 ", "表 ")):
            captions.append(text)
            continue
        sentence_groups.append(split_sentences(text))

    mandatory: list[str] = []
    seen: set[str] = set()

    def add(sentence: str) -> None:
        clean = normalize_paragraph(sentence)
        if clean and clean not in seen:
            seen.add(clean)
            mandatory.append(clean)

    for group in sentence_groups:
        if group:
            add(group[0])

    section_terms = PRIORITY_TERMS_BY_SECTION.get(section, [])
    terms = section_terms + [term for term in GLOBAL_PRIORITY_TERMS if term not in section_terms]
    for group in sentence_groups:
        for sentence in group[1:]:
            if any(term in sentence for term in terms):
                add(sentence)

    selected: list[str] = []
    total = 0
    for sentence in mandatory:
        if total + len(sentence) <= target + 120 or not selected:
            selected.append(sentence)
            total += len(sentence)

    for group in sentence_groups:
        for sentence in group:
            clean = normalize_paragraph(sentence)
            if clean in selected:
                continue
            if total + len(clean) > target:
                continue
            selected.append(clean)
            total += len(clean)

    if len("".join(selected)) > target + 160:
        trimmed: list[str] = []
        total = 0
        required_terms = section_terms or GLOBAL_PRIORITY_TERMS
        for sentence in selected:
            keep_for_term = any(term in sentence for term in required_terms)
            if total + len(sentence) <= target or keep_for_term:
                trimmed.append(sentence)
                total += len(sentence)
        selected = trimmed

    # Restore original order.
    order = {s: i for i, group in enumerate(sentence_groups) for s in map(normalize_paragraph, group)}
    selected = sorted(dict.fromkeys(selected), key=lambda s: order.get(s, 10_000))

    output = group_sentences(selected)
    if captions:
        # Keep figure/table placeholders, but avoid crowding very short sections.
        output.extend(captions[:2])
    return output


def collect_sections(doc: Document) -> dict[str, list[int]]:
    section_map: dict[str, list[int]] = {}
    current: str | None = None
    for index, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if is_heading(text):
            current = text
            section_map.setdefault(current, [])
            continue
        if current and text:
            section_map.setdefault(current, []).append(index)
    return section_map


def replace_section(doc: Document, heading: str, new_paragraphs: list[str], section_map: dict[str, list[int]]) -> None:
    indices = section_map.get(heading, [])
    if not indices:
        return
    if len(new_paragraphs) > len(indices):
        head = new_paragraphs[: len(indices) - 1]
        tail = "".join(new_paragraphs[len(indices) - 1 :])
        new_paragraphs = head + [tail]
    for offset, index in enumerate(indices[: len(new_paragraphs)]):
        doc.paragraphs[index].text = new_paragraphs[offset]
    for index in reversed(indices[len(new_paragraphs) :]):
        delete_paragraph(doc.paragraphs[index])


def main() -> None:
    if not SAFE_BASE.exists():
        raise FileNotFoundError(SAFE_BASE)
    if DOCX_PATH.exists() and not BACKUP_PATH.exists():
        shutil.copy2(DOCX_PATH, BACKUP_PATH)
    # Always begin from the verified readable 49k-character version.
    shutil.copy2(SAFE_BASE, DOCX_PATH)

    doc = Document(DOCX_PATH)

    report: list[str] = []
    for heading, target in SECTION_TARGETS.items():
        section_map = collect_sections(doc)
        indices = section_map.get(heading, [])
        current_paras = [doc.paragraphs[i].text.strip() for i in indices if doc.paragraphs[i].text.strip()]
        before = sum(len(p) for p in current_paras)
        if not current_paras or before <= target:
            continue
        new_paras = condense_text(heading, current_paras, target)
        replace_section(doc, heading, new_paras, section_map)
        after = sum(len(p) for p in new_paras)
        report.append(f"{heading}\t{before}->{after}")

    doc.save(DOCX_PATH)

    final_doc = Document(DOCX_PATH)
    paras = [p.text.strip() for p in final_doc.paragraphs if p.text.strip()]
    full = "\n".join(paras)
    report_path = Path("output/stage8-45page-rewrite-report.md")
    report_path.write_text(
        "# 45 页控制压缩报告\n\n"
        f"- 文件：`{DOCX_PATH.as_posix()}`\n"
        f"- 段落数：{len(paras)}\n"
        f"- 字符数：{sum(len(p) for p in paras)}\n"
        f"- 表格数：{len(final_doc.tables)}\n"
        f"- 问号数量：{full.count('?')}\n"
        f"- 本节关键信息摘要数量：{full.count('本节关键信息摘要')}\n\n"
        "## 压缩章节\n\n"
        + "\n".join(f"- {line}" for line in report)
        + "\n",
        encoding="utf-8",
    )
    print(report_path.as_posix())


if __name__ == "__main__":
    main()
