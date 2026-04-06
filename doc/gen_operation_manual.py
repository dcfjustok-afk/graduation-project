"""
生成《基于区块链的可信任务日志审计系统 · 操作手册》DOCX 文档。
用法: python doc/gen_operation_manual.py
输出: doc/操作手册.docx
"""

import os, io
from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# ── 路径 ──────────────────────────────────────────────
ROOT = Path(__file__).resolve().parent
SHOTS = ROOT / "screenshots"
OUT   = ROOT / "操作手册.docx"

# ── 标注工具 ──────────────────────────────────────────
def _font(size=18):
    """尝试加载中文字体"""
    candidates = [
        "C:/Windows/Fonts/msyh.ttc",   # 微软雅黑
        "C:/Windows/Fonts/simhei.ttf",  # 黑体
        "C:/Windows/Fonts/simsun.ttc",  # 宋体
    ]
    for f in candidates:
        if os.path.exists(f):
            try:
                return ImageFont.truetype(f, size)
            except Exception:
                continue
    return ImageFont.load_default()


def draw_arrow(draw, start, end, color="red", width=3):
    """在 PIL ImageDraw 上画带箭头的线段"""
    import math
    draw.line([start, end], fill=color, width=width)
    # 箭头
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    arrow_len = 16
    arrow_angle = math.pi / 6
    x1 = end[0] - arrow_len * math.cos(angle - arrow_angle)
    y1 = end[1] - arrow_len * math.sin(angle - arrow_angle)
    x2 = end[0] - arrow_len * math.cos(angle + arrow_angle)
    y2 = end[1] - arrow_len * math.sin(angle + arrow_angle)
    draw.polygon([end, (x1, y1), (x2, y2)], fill=color)


def draw_label(draw, pos, text, font, bg="red", fg="white"):
    """在指定位置画一个圆角标签"""
    bbox = font.getbbox(text)
    tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]
    pad = 6
    x, y = pos
    draw.rounded_rectangle(
        [x, y, x + tw + 2 * pad, y + th + 2 * pad],
        radius=6, fill=bg
    )
    draw.text((x + pad, y + pad - 2), text, fill=fg, font=font)


def draw_rect_outline(draw, box, color="red", width=3):
    """画矩形虚线框（高亮区域）"""
    draw.rectangle(box, outline=color, width=width)


def annotate_image(src_path, annotations, out_path=None):
    """
    对截图添加标注。
    annotations: list of dict, 每个 dict 可含:
      - type: "arrow" | "label" | "rect" | "arrow_label"
      - 对应参数
    返回标注后的 Image 或保存到 out_path。
    """
    img = Image.open(src_path).convert("RGB")
    draw = ImageDraw.Draw(img)
    font_sm = _font(20)
    font_lg = _font(26)

    for a in annotations:
        t = a["type"]
        if t == "arrow":
            draw_arrow(draw, a["start"], a["end"], a.get("color", "red"), a.get("width", 3))
        elif t == "label":
            draw_label(draw, a["pos"], a["text"], font_lg, a.get("bg", "#e53935"), a.get("fg", "white"))
        elif t == "rect":
            draw_rect_outline(draw, a["box"], a.get("color", "red"), a.get("width", 3))
        elif t == "arrow_label":
            draw_arrow(draw, a["start"], a["end"], a.get("color", "red"), a.get("width", 3))
            draw_label(draw, a["label_pos"], a["text"], font_sm, a.get("bg", "#e53935"), "white")

    if out_path:
        img.save(out_path, quality=92)
    return img


def img_to_bytes(img: Image.Image) -> io.BytesIO:
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    buf.seek(0)
    return buf


# ── DOCX 构建 ──────────────────────────────────────────
def set_cell_shading(cell, color_hex):
    """设置表格单元格背景色"""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:fill'): color_hex,
    })
    shading.append(shd)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h


def add_para(doc, text, bold=False, indent_cm=0, font_size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.bold = bold
    if indent_cm:
        p.paragraph_format.left_indent = Cm(indent_cm)
    return p


def add_image(doc, img_or_path, width_inches=6.2, caption=None):
    """插入图片（支持 PIL Image 或文件路径）"""
    if isinstance(img_or_path, Image.Image):
        buf = img_to_bytes(img_or_path)
        doc.add_picture(buf, width=Inches(width_inches))
    else:
        doc.add_picture(str(img_or_path), width=Inches(width_inches))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(caption)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run.italic = True


def add_step_block(doc, step_num, title, description, img, caption):
    """添加一个操作步骤块：标题 + 描述 + 带标注截图"""
    add_heading(doc, f"步骤 {step_num}：{title}", level=2)
    for line in description:
        add_para(doc, line, indent_cm=0.5)
    add_image(doc, img, caption=caption)
    doc.add_paragraph()  # 空行


# ── 主流程 ──────────────────────────────────────────
def build():
    doc = Document()

    # ── 样式设置 ──
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # ── 封面 ──
    doc.add_paragraph()
    doc.add_paragraph()
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run("基于区块链的可信任务日志审计系统")
    run.font.size = Pt(26)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_sub.add_run("系统操作手册")
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x42, 0x42, 0x42)

    doc.add_paragraph()
    doc.add_paragraph()

    # 封面信息表
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info = [
        ("文档版本", "V1.0"),
        ("编写日期", "2026年4月"),
        ("系统版本", "graduation-project 1.0.0"),
        ("适用对象", "系统操作人员 / 毕业答辩演示"),
    ]
    for i, (k, v) in enumerate(info):
        cell_k, cell_v = table.rows[i].cells
        cell_k.text = k
        cell_v.text = v
        for cell in (cell_k, cell_v):
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(11)

    doc.add_page_break()

    # ── 目录页 ──
    add_heading(doc, "目  录", level=1)
    toc_items = [
        "一、系统概述",
        "二、环境准备与启动",
        "三、系统总览页面",
        "四、日志生成操作",
        "    4.1 单条日志生成",
        "    4.2 批量日志生成",
        "五、日志中心",
        "六、审计操作",
        "    6.1 一键批量审计",
        "    6.2 审计结果查看",
        "七、异常告警",
        "八、完整演示流程总结",
    ]
    for item in toc_items:
        add_para(doc, item, font_size=12)
    doc.add_page_break()

    # ═══════════════════════════════════════════
    # 一、系统概述
    # ═══════════════════════════════════════════
    add_heading(doc, "一、系统概述", level=1)
    add_para(doc, "本系统是一个基于区块链的可信任务日志审计系统，实现了以下核心功能：")
    features = [
        "日志自动采集：Agent 增量采集日志文件，Web 端支持预设模板造数；",
        "链下存储：日志原文存入 SQLite 数据库，方便全文检索；",
        "链上存证：对日志内容计算 SHA-256 哈希后，通过 Solidity 智能合约写入 Hardhat 本地链；",
        "审计核验：重新计算哈希并与链上记录比对，哈希一致则「审计通过」，不一致则标记「异常」；",
        "异常告警：篡改检测失败时自动生成告警，按高危/中危/提示分级展示。",
    ]
    for f in features:
        add_para(doc, f"• {f}", indent_cm=0.8)

    doc.add_paragraph()
    add_para(doc, "系统架构示意：", bold=True)
    add_para(doc, "日志文件 → Agent 采集 → 后端 API → SQLite 存储 + 链上哈希存证 → 审计比对 → 可视化仪表盘", indent_cm=0.5)
    doc.add_paragraph()

    # ═══════════════════════════════════════════
    # 二、环境准备与启动
    # ═══════════════════════════════════════════
    add_heading(doc, "二、环境准备与启动", level=1)
    add_para(doc, "系统运行需要以下环境：", bold=True)
    envs = [
        "Node.js ≥ 18 （推荐 20 LTS）",
        "npm ≥ 9",
        "Python 3.x （仅生成本文档时需要）",
    ]
    for e in envs:
        add_para(doc, f"• {e}", indent_cm=0.8)

    doc.add_paragraph()
    add_para(doc, "启动步骤：", bold=True)
    steps = [
        ("1. 安装依赖", "npm install"),
        ("2. 启动 Hardhat 本地链", "npx --prefix packages/contracts hardhat node"),
        ("3. 部署智能合约", "npm run chain:deploy"),
        ("4. 启动后端服务", "npm --prefix apps/server run dev"),
        ("5. 启动前端", "npm --prefix apps/web run dev"),
        ("6. 访问系统", "浏览器打开 http://localhost:5173"),
    ]
    for label, cmd in steps:
        p = doc.add_paragraph()
        run_b = p.add_run(f"{label}：")
        run_b.bold = True
        run_b.font.size = Pt(11)
        run_c = p.add_run(f"  {cmd}")
        run_c.font.size = Pt(10)
        run_c.font.color.rgb = RGBColor(0x00, 0x7A, 0xCC)

    add_para(doc, "")
    add_para(doc, "提示：也可使用一键启动命令 npm run dev 同时启动所有服务。", font_size=10)
    doc.add_page_break()

    # ═══════════════════════════════════════════
    # 三、系统总览页面
    # ═══════════════════════════════════════════
    add_heading(doc, "三、系统总览页面", level=1)
    add_para(doc, "登录系统后，默认进入「系统总览」页面，展示系统整体运行状态。")
    doc.add_paragraph()

    # 标注 dashboard 截图
    img1 = annotate_image(SHOTS / "09_dashboard_with_data.png", [
        {"type": "rect", "box": (0, 0, 200, 900), "color": "#1677ff", "width": 4},
        {"type": "arrow_label", "start": (220, 60), "end": (200, 60), "label_pos": (225, 45), "text": "① 侧边导航栏", "color": "#1677ff"},
        {"type": "rect", "box": (200, 0, 1440, 80), "color": "#52c41a", "width": 3},
        {"type": "arrow_label", "start": (820, 100), "end": (820, 78), "label_pos": (700, 102), "text": "② 顶部状态栏", "color": "#52c41a"},
        {"type": "rect", "box": (205, 300, 1435, 445), "color": "#e53935", "width": 3},
        {"type": "arrow_label", "start": (1200, 280), "end": (1200, 302), "label_pos": (1050, 255), "text": "③ 统计卡片区", "color": "#e53935"},
        {"type": "rect", "box": (205, 450, 900, 650), "color": "#ff9800", "width": 3},
        {"type": "arrow_label", "start": (550, 670), "end": (550, 648), "label_pos": (420, 672), "text": "④ 审计时间线", "color": "#ff9800"},
    ])
    add_image(doc, img1, caption="图 3-1  系统总览页面（含数据状态）")
    doc.add_paragraph()

    add_para(doc, "页面布局说明：", bold=True)
    layout_desc = [
        "① 侧边导航栏：包含 5 个功能模块入口 —— 系统总览、日志生成台、日志中心、审计中心、异常告警。底部显示当前数据模式（真实数据 / 链上存证）。",
        "② 顶部状态栏：显示系统名称、核心流程链路文字，以及系统运行状态标签（绿色「系统运行中」）和区块链网络标签（蓝色「Hardhat 本地链」）。",
        "③ 统计卡片区：四张指标卡片分别展示「日志总量」「链上存证记录」「审计记录数」「活动告警数」，数值实时从后端 API 读取。",
        "④ 审计时间线：按时间倒序展示最近的审计事件，显示审计结果摘要信息。",
    ]
    for d in layout_desc:
        add_para(doc, d, indent_cm=0.5, font_size=10)

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # 四、日志生成操作
    # ═══════════════════════════════════════════
    add_heading(doc, "四、日志生成操作", level=1)
    add_para(doc, "日志生成台提供可视化造数功能，支持预设模板快速填充和自定义表单输入。")
    doc.add_paragraph()

    # 4.1 单条日志
    add_heading(doc, "4.1 单条日志生成", level=2)
    add_para(doc, "操作流程：点击预设模板 → 调整表单 → 点击「生成单条日志」按钮。")
    doc.add_paragraph()

    img2 = annotate_image(SHOTS / "02_log_generator.png", [
        {"type": "rect", "box": (210, 190, 1430, 290), "color": "#1677ff", "width": 3},
        {"type": "arrow_label", "start": (1200, 170), "end": (1200, 192), "label_pos": (1040, 145), "text": "① 预设模板区", "color": "#1677ff"},
        {"type": "rect", "box": (210, 290, 900, 750), "color": "#52c41a", "width": 3},
        {"type": "arrow_label", "start": (180, 520), "end": (212, 520), "label_pos": (15, 505), "text": "② 表单输入区", "color": "#52c41a"},
        {"type": "rect", "box": (210, 750, 900, 830), "color": "#e53935", "width": 3},
        {"type": "arrow_label", "start": (550, 850), "end": (550, 828), "label_pos": (400, 852), "text": "③ 操作按钮", "color": "#e53935"},
        {"type": "rect", "box": (910, 290, 1430, 750), "color": "#ff9800", "width": 3},
        {"type": "arrow_label", "start": (1450, 520), "end": (1428, 520), "label_pos": (1455, 505), "text": "④ 预览面板", "color": "#ff9800"},
    ])
    add_image(doc, img2, caption="图 4-1  日志生成台页面布局")
    doc.add_paragraph()

    add_para(doc, "区域说明：", bold=True)
    gen_desc = [
        "① 预设模板区：提供三种典型场景模板 —— 「日报任务/INFO」（蓝色）、「备份延迟/WARN」（金色）、「权限异常/ERROR」（红色）。点击模板自动填充表单。",
        "② 表单输入区：包含任务名称、来源类型、来源路径、日志级别、日志内容等字段。带 * 号为必填项。",
        "③ 操作按钮：「生成单条日志」提交 1 条；「批量生成日志」根据数量设置提交多条；「重置表单」清空所有字段。",
        "④ 预览面板：实时预览当前级别、批量数量、任务信息，以及提交结果反馈。",
    ]
    for d in gen_desc:
        add_para(doc, d, indent_cm=0.5, font_size=10)
    doc.add_paragraph()

    # 单条提交结果
    img3 = annotate_image(SHOTS / "03_log_generator_single_result.png", [
        {"type": "rect", "box": (910, 680, 1430, 780), "color": "#e53935", "width": 4},
        {"type": "arrow_label", "start": (1100, 660), "end": (1100, 682), "label_pos": (940, 635), "text": "提交成功，返回日志 ID", "color": "#e53935"},
    ])
    add_image(doc, img3, caption="图 4-2  单条日志提交成功结果")
    add_para(doc, "提交成功后，右侧提交结果面板会显示「最近创建日志 ID」，该日志已同步写入数据库并上链存证。", indent_cm=0.5, font_size=10)
    doc.add_paragraph()

    # 4.2 批量日志
    add_heading(doc, "4.2 批量日志生成", level=2)
    add_para(doc, "切换到「备份延迟/WARN」模板后，设置生成数量（如 3 条），点击「批量生成日志」。")
    doc.add_paragraph()

    img4 = annotate_image(SHOTS / "04_log_generator_batch_result.png", [
        {"type": "rect", "box": (910, 640, 1430, 850), "color": "#e53935", "width": 4},
        {"type": "arrow_label", "start": (1170, 620), "end": (1170, 642), "label_pos": (1010, 595), "text": "批量提交结果统计", "color": "#e53935"},
    ])
    add_image(doc, img4, caption="图 4-3  批量日志生成结果（成功 3 条，失败 0 条）")
    add_para(doc, "批量生成完成后，提交结果区域会显示成功/失败条数统计。所有成功日志均已自动完成 SHA-256 哈希计算并上链。", indent_cm=0.5, font_size=10)

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # 五、日志中心
    # ═══════════════════════════════════════════
    add_heading(doc, "五、日志中心", level=1)
    add_para(doc, "日志中心以表格形式展示所有已采集的日志记录，支持关键字搜索和分页浏览。")
    doc.add_paragraph()

    img5 = annotate_image(SHOTS / "05_logs_center.png", [
        {"type": "rect", "box": (210, 80, 1435, 145), "color": "#1677ff", "width": 3},
        {"type": "arrow_label", "start": (1200, 60), "end": (1200, 82), "label_pos": (1050, 35), "text": "① 搜索 & 刷新", "color": "#1677ff"},
        {"type": "rect", "box": (210, 145, 1435, 190), "color": "#52c41a", "width": 3},
        {"type": "arrow_label", "start": (180, 167), "end": (212, 167), "label_pos": (15, 152), "text": "② 表头列", "color": "#52c41a"},
        {"type": "rect", "box": (520, 195, 620, 500), "color": "#ff9800", "width": 3},
        {"type": "arrow_label", "start": (570, 510), "end": (570, 498), "label_pos": (470, 515), "text": "③ 日志级别", "color": "#ff9800"},
        {"type": "rect", "box": (640, 195, 760, 500), "color": "#e53935", "width": 3},
        {"type": "arrow_label", "start": (700, 510), "end": (700, 498), "label_pos": (620, 515), "text": "④ 状态标签", "color": "#e53935"},
    ])
    add_image(doc, img5, caption="图 5-1  日志中心页面")
    doc.add_paragraph()

    add_para(doc, "页面功能说明：", bold=True)
    log_desc = [
        "① 搜索与刷新：输入框支持按日志编号、任务名、来源文件、哈希摘要等关键字搜索；刷新按钮重新加载数据。",
        "② 表头：共 8 列 —— 日志编号、任务名称、来源文件、级别、状态、提交时间、哈希摘要、审计说明。",
        "③ 日志级别：以彩色标签展示 —— INFO（蓝）、WARN（金）、ERROR（红），直观区分日志严重程度。",
        "④ 状态标签：绿色「已上链」表示已完成区块链存证；审计后会变为「审计通过」或「异常」。",
    ]
    for d in log_desc:
        add_para(doc, d, indent_cm=0.5, font_size=10)

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # 六、审计操作
    # ═══════════════════════════════════════════
    add_heading(doc, "六、审计操作", level=1)
    add_para(doc, "审计中心是系统核心功能页面，负责执行链上哈希比对审计，检测日志数据是否被篡改。")
    doc.add_paragraph()

    # 6.1 审计前
    add_heading(doc, "6.1 一键批量审计", level=2)
    add_para(doc, "进入审计中心后，点击页面顶部的「一键批量审计」按钮，系统将自动对所有日志执行审计。")
    doc.add_paragraph()

    img6 = annotate_image(SHOTS / "06_audit_before.png", [
        {"type": "rect", "box": (470, 82, 760, 120), "color": "#e53935", "width": 4},
        {"type": "arrow_label", "start": (615, 130), "end": (615, 118), "label_pos": (490, 133), "text": "点击此按钮开始审计", "color": "#e53935"},
        {"type": "rect", "box": (210, 125, 1435, 210), "color": "#ff9800", "width": 3},
        {"type": "arrow_label", "start": (1200, 225), "end": (1200, 208), "label_pos": (1050, 228), "text": "审计统计指标", "color": "#ff9800"},
    ])
    add_image(doc, img6, caption="图 6-1  审计中心（审计执行前）")
    doc.add_paragraph()

    add_para(doc, "审计流程原理：", bold=True)
    audit_desc = [
        "1. 从数据库读取每条日志的原始内容，重新计算 SHA-256 哈希值；",
        "2. 将重新计算的哈希与数据库中存储的提交时哈希进行比对（链下校验）；",
        "3. 从区块链智能合约读取该日志的链上哈希记录，与本地重算哈希比对（链上校验）；",
        "4. 若两次比对均一致，标记「审计通过」；若不一致，标记「异常」并自动生成告警。",
    ]
    for d in audit_desc:
        add_para(doc, d, indent_cm=0.5, font_size=10)
    doc.add_paragraph()

    # 6.2 审计后
    add_heading(doc, "6.2 审计结果查看", level=2)
    add_para(doc, "审计完成后，页面自动刷新，展示审计结果统计和各日志的审计状态。")
    doc.add_paragraph()

    img7 = annotate_image(SHOTS / "07_audit_after.png", [
        {"type": "rect", "box": (210, 125, 1435, 210), "color": "#52c41a", "width": 4},
        {"type": "arrow_label", "start": (400, 225), "end": (400, 208), "label_pos": (260, 228), "text": "统计数值已更新", "color": "#52c41a"},
        {"type": "rect", "box": (210, 210, 900, 480), "color": "#1677ff", "width": 3},
        {"type": "arrow_label", "start": (550, 500), "end": (550, 478), "label_pos": (400, 502), "text": "审计执行面板 & 时间线", "color": "#1677ff"},
        {"type": "rect", "box": (900, 210, 1435, 480), "color": "#ff9800", "width": 3},
        {"type": "arrow_label", "start": (1170, 500), "end": (1170, 478), "label_pos": (1020, 502), "text": "状态分布 & 趋势图", "color": "#ff9800"},
    ])
    add_image(doc, img7, caption="图 6-2  审计完成后页面状态")
    add_para(doc, "审计完成后，「待处理」指标更新为实际数值，状态统计图饼图和日志趋势折线图同步刷新。每条日志的审计记录会显示在时间线和底部表格中。", indent_cm=0.5, font_size=10)

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # 七、异常告警
    # ═══════════════════════════════════════════
    add_heading(doc, "七、异常告警", level=1)
    add_para(doc, "当审计检测到日志被篡改（哈希不匹配）时，系统自动生成告警记录，在异常告警页面集中展示。")
    doc.add_paragraph()

    img8 = annotate_image(SHOTS / "08_alerts_empty.png", [
        {"type": "rect", "box": (210, 80, 1435, 340), "color": "#1677ff", "width": 3},
        {"type": "arrow_label", "start": (500, 340), "end": (500, 338), "label_pos": (340, 348), "text": "异常分布 & 关注重点", "color": "#1677ff"},
        {"type": "rect", "box": (210, 350, 1435, 650), "color": "#ff9800", "width": 3},
        {"type": "arrow_label", "start": (820, 660), "end": (820, 648), "label_pos": (700, 662), "text": "告警卡片列表区（当前为空）", "color": "#ff9800"},
    ])
    add_image(doc, img8, caption="图 7-1  异常告警页面（暂无告警）")
    doc.add_paragraph()

    add_para(doc, "告警分级说明：", bold=True)
    alert_desc = [
        "• 高危（红色）：哈希比对完全不一致，确认数据被篡改。",
        "• 中危（金色）：部分校验异常，可能存在数据不完整。",
        "• 提示（蓝色）：链上记录暂不可用等轻微异常。",
    ]
    for d in alert_desc:
        add_para(doc, d, indent_cm=0.5, font_size=10)
    doc.add_paragraph()

    add_para(doc, "篡改检测演示：", bold=True)
    add_para(doc, "可通过终端运行命令 npm --prefix apps/server run experiment:tamper 执行自动化篡改实验。该脚本会：", indent_cm=0.5, font_size=10)
    tamper_steps = [
        "1. 提交一条正常日志并上链；",
        "2. 直接修改数据库中该日志的 content 字段（模拟篡改）；",
        "3. 重新执行审计；",
        "4. 审计结果显示「异常」，并在异常告警页面生成高危告警卡片。",
    ]
    for s in tamper_steps:
        add_para(doc, s, indent_cm=1.0, font_size=10)

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # 八、完整演示流程总结
    # ═══════════════════════════════════════════
    add_heading(doc, "八、完整演示流程总结", level=1)
    add_para(doc, "以下表格汇总了完整的系统操作演示流程：")
    doc.add_paragraph()

    # 流程表格
    table = doc.add_table(rows=9, cols=3)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["步骤", "操作", "预期结果"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    flow_data = [
        ("1", "进入「系统总览」", "确认系统运行中，数据来源显示「真实后端」"),
        ("2", "进入「日志生成台」→ 选择 INFO 模板 → 生成单条日志", "提交成功，返回日志 ID"),
        ("3", "选择 WARN 模板 → 批量生成 3 条日志", "显示成功 3 条，失败 0 条"),
        ("4", "选择 ERROR 模板 → 生成 1 条日志", "提交成功，共计 5 条日志"),
        ("5", "进入「日志中心」", "表格展示 5 条日志，全部状态为「已上链」"),
        ("6", "进入「审计中心」→ 点击「一键批量审计」", "统计数值更新，日志状态变更"),
        ("7", "终端运行篡改实验脚本", "日志被篡改后审计标记为「异常」"),
        ("8", "进入「异常告警」", "看到高危告警卡片"),
    ]
    for i, (step, op, result) in enumerate(flow_data):
        row = table.rows[i + 1]
        row.cells[0].text = step
        row.cells[1].text = op
        row.cells[2].text = result

    doc.add_paragraph()
    doc.add_paragraph()

    # 技术指标
    add_heading(doc, "关键技术指标", level=2)

    t2 = doc.add_table(rows=5, cols=2)
    t2.style = 'Light Grid Accent 1'
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    t2.rows[0].cells[0].text = "指标"
    t2.rows[0].cells[1].text = "数值"
    for run in t2.rows[0].cells[0].paragraphs[0].runs:
        run.bold = True
    for run in t2.rows[0].cells[1].paragraphs[0].runs:
        run.bold = True

    metrics = [
        ("日志提交平均耗时（含链上写入）", "~170 ms"),
        ("审计吞吐量", "~5.7 req/s"),
        ("篡改检测准确率", "100%（SHA-256 + 链上比对）"),
        ("智能合约事件覆盖", "日志写入即触发 LogStored 事件"),
    ]
    for i, (k, v) in enumerate(metrics):
        t2.rows[i + 1].cells[0].text = k
        t2.rows[i + 1].cells[1].text = v

    doc.add_paragraph()
    doc.add_paragraph()

    # 结束语
    p_end = doc.add_paragraph()
    p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_end.add_run("—— 文档结束 ——")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # ── 保存 ──
    doc.save(str(OUT))
    print(f"✅ 操作手册已生成: {OUT}")


if __name__ == "__main__":
    build()
